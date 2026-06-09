#!/usr/bin/env python3
"""
Generate the final WASABI / HumanEnerDIA / OVOS-EnMS DOCX package.

This script intentionally reads no private .env file and does not modify
application source code. It produces delivery documents, source Markdown, and
diagram assets under docs/final-delivery/.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "docs" / "final-delivery"
ASSET_DIR = OUT_DIR / "assets"
SOURCE_DIR = OUT_DIR / "source"
DOC_DATE = date(2026, 6, 9).isoformat()
PROJECT_NAME = "WASABI / HumanEnerDIA / OVOS-EnMS"
DOC_VERSION = "1.1"
DOC_STATUS = "Final stakeholder-ready documentation package"


TextBlock = Union[str, Tuple[str, str]]


@dataclass
class Section:
    title: str
    paragraphs: List[TextBlock] = field(default_factory=list)
    bullets: List[str] = field(default_factory=list)
    table: Optional[Tuple[List[str], List[List[str]]]] = None
    table_followup: List[TextBlock] = field(default_factory=list)
    figure: Optional[Tuple[str, str]] = None
    subsections: List["Section"] = field(default_factory=list)


@dataclass
class DocSpec:
    filename: str
    title: str
    purpose: str
    audience: str
    evidence_note: str
    sections: List[Section]


def safe_path(path: Path) -> str:
    return str(path)


EVIDENCE = {
    "compose": "docker-compose.yml",
    "setup": "setup.sh",
    "verifier": "verify.sh",
    "readme": "README.md",
    "nginx": "nginx/nginx.conf; nginx/conf.d/default.conf",
    "database_schema": "database/init/02-schema.sql; database/init/03-timescaledb-setup.sql; database/init/04-functions.sql",
    "seed_data": "database/init/06-seed-data.sql",
    "analytics_main": "analytics/main.py",
    "analytics_routes": "analytics/api/routes/",
    "analytics_services": "analytics/services/",
    "analytics_models": "analytics/models/",
    "reports": "analytics/api/routes/reports.py; analytics/reports/; analytics/reports_v2/",
    "grafana": "grafana/provisioning/; grafana/dashboards/",
    "nodered": "nodered/data/flows.json; nodered/settings.js; nodered/package.json",
    "simulator": "simulator/main.py; simulator/api/routes.py; simulator/simulator_manager.py; simulator/mqtt_publisher.py",
    "auth": "auth-service/app.py; auth-service/auth_service.py; database/init/05-auth-schema.sql",
    "chatbot": "chatbot/server/index.js; chatbot/rasa/actions/actions.py; chatbot/rasa/qa_data.json",
    "ovos_readme": "OVOS-EnMS repository: README.md; enms-ovos-skill/README.md",
    "ovos_bridge": "OVOS-EnMS repository: enms-ovos-skill/bridge/ovos_rest_bridge.py",
    "ovos_skill": "OVOS-EnMS repository: enms-ovos-skill/enms_ovos_skill/__init__.py",
    "ovos_parser": "OVOS-EnMS repository: enms-ovos-skill/enms_ovos_skill/lib/intent_parser.py; lib/adapt_parser.py; lib/llm_parser.py",
    "ovos_validator": "OVOS-EnMS repository: enms-ovos-skill/enms_ovos_skill/lib/validator.py",
    "ovos_client": "OVOS-EnMS repository: enms-ovos-skill/enms_ovos_skill/lib/api_client.py",
    "ovos_formatter": "OVOS-EnMS repository: enms-ovos-skill/enms_ovos_skill/lib/response_formatter.py",
    "ovos_config": "OVOS-EnMS repository: docker-compose.yml; Dockerfile; enms-ovos-skill/config.yaml.template; enms-ovos-skill/settings.docker.json; enms-ovos-skill/settingsmeta.yaml",
}


SERVICE_ROWS = [
    ["nginx", "nginx:1.25-alpine", "8080; 8443 mapped but HTTPS server block is optional/commented until certificates are configured", "Public gateway, portal/static hosting, reverse proxy", "Yes"],
    ["postgres", "timescale/timescaledb:latest-pg16", "5433", "PostgreSQL with TimescaleDB extension and persistent data volume", "Yes"],
    ["mqtt", "build ./mqtt", "1883, 9001", "Mosquitto telemetry broker with configured credentials", "Yes"],
    ["redis", "redis:7-alpine", "6380", "Redis cache and Pub/Sub support for analytics event paths", "Yes"],
    ["simulator", "build ./simulator", "internal 8003", "FastAPI synthetic telemetry generator loaded from database machines", "Yes"],
    ["nodered", "build ./nodered", "1881", "MQTT-to-database ingestion and automation flow runtime", "Yes"],
    ["grafana", "grafana/grafana:10.2.0", "3001, /grafana", "Provisioned dashboards backed by PostgreSQL/TimescaleDB", "Yes"],
    ["analytics", "build ./analytics", "8001, /api/analytics", "FastAPI analytics, KPI, reports, ISO 50001, and OVOS proxy APIs", "Yes"],
    ["auth-service", "build ./auth-service", "5500", "Flask auth, admin, contact, pilot/application APIs", "Yes"],
    ["rasa-actions", "build ./chatbot/rasa", "5055", "Rasa custom action server", "Yes"],
    ["rasa", "build ./chatbot/rasa", "5005", "Rasa NLU text chatbot server", "Yes"],
    ["chatbot", "build ./chatbot", "5006", "Express backend and built chatbot frontend proxying to Rasa and OVOS", "Yes"],
]


ACCESS_ROWS = [
    ["Unified portal", "http://<host>:8080", "Served by Nginx from portal/public"],
    ["Grafana", "http://<host>:8080/grafana", "Sub-path proxy to Grafana with provisioned dashboards"],
    ["Analytics UI", "http://<host>:8080/analytics/ui/", "FastAPI-rendered analytics templates"],
    ["Analytics API docs", "http://<host>:8080/api/analytics/docs", "Nginx proxy to analytics OpenAPI docs"],
    ["Simulator docs", "http://<host>:8080/api/simulator/docs", "Nginx proxy to simulator OpenAPI docs"],
    ["Node-RED", "http://<host>:1881 or http://<host>:8080/nodered/", "Admin UI protected by Node-RED credentials"],
    ["OVOS bridge", "http://<host>:5000/health", "Available when the separate OVOS-EnMS runtime is deployed"],
]


API_ROUTE_ROWS = [
    ["Health and system", "/api/v1/health, /api/v1/stats/system, /api/v1/stats/connections", "analytics/main.py"],
    ["Baselines", "/api/v1/baseline/train, /deviation, /predict, /models, /train-seu", "analytics/api/routes/baseline.py"],
    ["KPIs", "/api/v1/kpi/sec, /peak-demand, /load-factor, /energy-cost, /carbon, /all, /factory", "analytics/api/routes/kpi.py"],
    ["Forecasting", "/api/v1/forecast/train/arima, /train/prophet, /predict, /demand, /peak, /short-term", "analytics/api/routes/forecast.py"],
    ["Anomalies", "/api/v1/anomaly/create, /detect, /search, /recent, /active, /{id}/resolve", "analytics/api/routes/anomaly.py"],
    ["Machines and time series", "/api/v1/machines, /machines/status/{name}, /timeseries/energy, /power, /latest/{id}", "analytics/api/routes/machines.py; timeseries.py"],
    ["ISO 50001 and SEUs", "/api/v1/iso50001/*, /api/v1/seus, /api/v1/reports/seu-performance", "analytics/api/routes/iso50001.py; seu.py; seus.py"],
    ["Reports", "/api/v1/reports/types, /generate, /preview, /v2/generate, /v2/download/{id}, /v2/status", "analytics/api/routes/reports.py"],
    ["OVOS integration", "/api/v1/ovos/*, /api/v1/ovos/voice/query, /voice/health, /voice/config", "analytics/api/routes/ovos.py; ovos_voice.py"],
    ["Visualization data", "/api/v1/sankey/data, /heatmap/hourly, /comparison/machines, /compare/machines", "analytics/api/routes/sankey.py; heatmap.py; comparison.py; compare.py"],
]


KPI_ROWS = [
    ["Specific Energy Consumption", "SEC = total energy kWh / total production units", "calculate_sec() over energy_readings_1hour and production_data_1hour", "database/init/04-functions.sql; /api/v1/kpi/sec"],
    ["Peak demand", "Maximum 15-minute peak_demand_kw in selected period", "calculate_peak_demand() over energy_readings_15min", "database/init/04-functions.sql; /api/v1/kpi/peak-demand"],
    ["Load factor", "Average power divided by maximum power", "calculate_load_factor() over energy_readings_15min", "database/init/04-functions.sql; /api/v1/kpi/load-factor"],
    ["Energy cost", "Energy multiplied by tariff rate; active time-of-use tariff selected when configured", "calculate_energy_cost() queries energy_tariffs with default fallback rate", "database/init/04-functions.sql; /api/v1/kpi/energy-cost"],
    ["Carbon intensity/emissions", "Energy multiplied by active carbon factor, with default factor fallback", "calculate_carbon_intensity() queries carbon_factors", "database/init/04-functions.sql; /api/v1/kpi/carbon"],
    ["Combined KPI response", "Aggregates SEC, peak demand, load factor, cost, and carbon", "calculate_all_kpis() and KPIService.calculate_all_kpis()", "database/init/04-functions.sql; analytics/services/kpi_service.py"],
]


DASHBOARD_ROWS = [
    ["SOTA Factory Overview", "Active machines, energy today, cost today, active anomalies, current power, machine status"],
    ["SOTA Machine Health", "Health score, current power, baseline variance, production, actual vs baseline, anomalies"],
    ["SOTA ISO 50001 EnPI", "EnPI score, energy savings, compliance rate, CUSUM, baseline vs actual, SEU performance"],
    ["SOTA Energy Cost Analytics", "Cost trend, time-of-use cost, top cost contributors, savings opportunities"],
    ["SOTA Environmental Impact", "Monthly carbon footprint, CO2 trend, emission intensity, emissions by machine"],
    ["SOTA Predictive Analytics", "Forecast metrics, forecast vs actual, accuracy trends, recent forecasts"],
    ["SOTA Anomaly Detection", "Active and critical anomalies, severity distribution, machine-hour heatmap, unresolved list"],
    ["SOTA ML Model Performance", "Active models, R2/RMSE, model performance trends, training history"],
    ["SOTA Operational Efficiency", "OEE, availability, performance rate, production vs energy efficiency"],
    ["SOTA Real-Time Production", "Live factory status, active machines, current power"],
    ["SOTA Executive Summary", "Operational concerns, 12-month energy trend, energy intensity, monthly summary"],
]


INTENT_ROWS = [
    ["energy_query", "Energy use questions by machine or factory scope"],
    ["power_query", "Current or historical power demand questions"],
    ["machine_status", "Machine running/offline/status checks"],
    ["factory_overview", "Factory/facility summaries, machine lists, aggregate status"],
    ["comparison", "Machine-to-machine comparisons"],
    ["ranking", "Top or lowest machines by energy, power, cost, efficiency, or alerts"],
    ["anomaly_detection", "Active/recent anomaly and alert queries"],
    ["cost_analysis", "Cost and spending questions"],
    ["forecast", "Forecasted demand and future energy usage"],
    ["baseline, baseline_models, baseline_explanation", "Baseline prediction, model inventory, and driver explanation"],
    ["driver_analysis", "Energy driver analysis for factory or SEU/machine context"],
    ["seus", "Significant Energy Use listing and context"],
    ["kpi, performance, production", "KPIs, performance analysis, production/OEE-related queries"],
    ["report", "Report type, preview, and generation workflows"],
    ["help, health", "Capability help and system health checks"],
]


LIMITATION_ROWS = [
    ["Runtime verification", "Compose validation is confirmed where stated. Live health checks are deployment-specific and require a running target environment."],
    ["OVOS deployment boundary", "The GitHub production base docker-compose.yml does not define an OVOS service. OVOS-EnMS is documented as a separate source repository and companion assistant runtime."],
    ["OVOS optional LLM fallback", "The OVOS-EnMS Dockerfile installs LLM fallback dependencies only when INSTALL_LLM_FALLBACK=true. Model availability must be verified in the OVOS-EnMS repository/runtime."],
    ["Third-party EnMS support", "OVOS portability is through a HumanEnerDIA-compatible API or adapter/proxy, not zero-code support for arbitrary vendor APIs."],
    ["Reports V2", "V2 report code is implemented, but some service calculations use derived, proportional, or placeholder values. Formal audit use requires independent validation of formulas, source data, tariff factors, carbon factors, and generated report semantics."],
    ["Simulator inventory", "The simulator code supports boiler in addition to compressor, HVAC, motor, pump, and injection molding. One simulator info response still lists five machine types."],
    ["Security posture", "The codebase provides secret placeholders, generated first-run credentials, JWT/bcrypt auth, health checks, and hardening guidance. Public production exposure still requires operator DNS/TLS/firewall/credential work."],
]


COMPONENT_RESPONSIBILITY_ROWS = [
    ["Nginx", "Public HTTP gateway, static portal host, reverse proxy, and health endpoint.", "Configured service and routing rules in docker-compose.yml and nginx/conf.d/default.conf."],
    ["Analytics", "Primary EnMS domain API: KPIs, baselines, forecasts, anomalies, reports, ISO 50001, machine/time-series data, OVOS proxy paths.", "FastAPI application and router registrations in analytics/main.py."],
    ["PostgreSQL/TimescaleDB", "Persistent relational and time-series storage, hypertables, continuous aggregates, SQL KPI functions, seed data.", "database/init/*.sql."],
    ["MQTT", "Telemetry broker for simulator/device payloads and Node-RED ingestion.", "mqtt service in docker-compose.yml and simulator MQTT publisher."],
    ["Node-RED", "MQTT topic parsing, payload validation, routing by data type, PostgreSQL writes, and simple ingestion monitoring.", "nodered/data/flows.json."],
    ["Grafana", "Provisioned dashboard runtime for operational, cost, carbon, EnPI, anomaly, forecast, model, and executive views.", "grafana/provisioning and grafana/dashboards JSON."],
    ["Simulator", "Synthetic factory telemetry service with configurable machine simulators and anomaly injection.", "simulator/main.py; simulator/simulator_manager.py; simulator/api/routes.py."],
    ["auth-service", "Registration, login, JWT verification, email verification/reset, admin APIs, contact and pilot factory forms.", "auth-service/app.py; auth-service/auth_service.py."],
    ["chatbot/Rasa", "Text help chatbot and custom Rasa action backed by qa_data.json.", "chatbot/server/index.js; chatbot/rasa/actions/actions.py."],
    ["OVOS-EnMS", "Separate natural-language/voice assistant runtime that calls a HumanEnerDIA-compatible API.", "OVOS-EnMS repository: bridge, skill, parser, validator, API client, formatter."],
]


DEPENDENCY_FLOW_ROWS = [
    ["Telemetry path", "simulator or external devices -> MQTT -> Node-RED -> PostgreSQL/TimescaleDB -> analytics/Grafana/reports/assistants"],
    ["Dashboard path", "browser -> Nginx -> Grafana -> PostgreSQL/TimescaleDB datasource"],
    ["Analytics path", "browser/API client -> Nginx -> analytics -> PostgreSQL/TimescaleDB and optional Redis"],
    ["Authentication path", "portal/auth pages -> Nginx -> auth-service -> PostgreSQL auth tables; SMTP is optional/configured through environment"],
    ["Rasa help path", "portal chatbot -> chatbot Express backend -> Rasa server -> Rasa action server -> qa_data.json"],
    ["OVOS operational path", "REST bridge or portal proxy -> OVOS messagebus -> EnMS skill -> HumanEnerDIA-compatible analytics API -> structured response"],
]


DEPLOYMENT_VARIANT_ROWS = [
    ["Base production Compose", "docker-compose.yml in GitHub production", "HumanEnerDIA services only. No OVOS service appears in the current production service list."],
    ["Separate OVOS-EnMS runtime", "OVOS-EnMS repository docker-compose.yml", "Companion assistant runtime exposing bridge/messagebus ports and connecting to a HumanEnerDIA-compatible API."],
    ["Evaluation/demo deployment", "setup.sh with generated first-run secrets and simulator auto-start", "Suitable for review and demonstration after health checks pass; production hardening remains a separate deployment responsibility."],
    ["Production-hardened deployment", "Operator-controlled DNS/TLS/firewall/backups/monitoring", "Supported by configuration hooks but not automatically completed by the repository."],
]


OPERATIONAL_RISK_ROWS = [
    ["Directly mapped internal ports", "PostgreSQL, MQTT, Redis, Grafana, Node-RED, analytics, auth, Rasa, chatbot, simulator ports are mapped for access/operations.", "Restrict with firewall or upstream network policy before public exposure."],
    ["Runtime data persistence", "Named Docker volumes retain database, Grafana, MQTT, Redis, and Node-RED data.", "Back up before destructive redeployments or volume removal."],
    ["Credentials", ".env is generated locally and ignored; .env.example contains placeholders.", "Rotate generated credentials and never disclose .env values."],
    ["Report semantics", "Some V2 report calculations use estimates/placeholders.", "Formal audit use requires independent formula, source-data, tariff-factor, and carbon-factor validation."],
    ["OVOS availability", "Production base Compose does not start OVOS.", "Deploy and verify OVOS-EnMS separately when assistant access is in scope."],
]


API_DESIGN_DETAIL_ROWS = [
    ["Route organization", "analytics/main.py mounts route modules under settings.API_PREFIX, normally /api/v1.", "FastAPI routers keep domain areas separate and expose OpenAPI docs."],
    ["Request/response models", "Several routes use Pydantic response models, especially forecast, heatmap, ISO 50001, performance, model-performance, voice proxy, and reports.", "Validation and schema visibility are strongest in typed routes."],
    ["Database access", "Routes and services use the analytics database module and async connection pool where implemented.", "Data-heavy routes prefer SQL queries and database functions over in-memory mock state."],
    ["Legacy/deprecated routes", "OVOS and OVOS training routes remain mounted with comments noting newer factory/analytics/baseline paths.", "Document as compatibility surface, not a separate service."],
    ["UI route surface", "analytics/api/routes/ui_routes.py serves analytics UI pages for dashboard/baseline/anomaly/KPI/forecast/Sankey/heatmap/comparison/model-performance paths.", "These are FastAPI-rendered pages, distinct from Grafana dashboards."],
]


DATABASE_DESIGN_ROWS = [
    ["Core dimensions", "factories, machines", "Facility and machine metadata, including machine type, factory association, active flag, rated power, and topic-related context."],
    ["Telemetry facts", "energy_readings, production_data, environmental_data, machine_status", "High-frequency operational data and current machine state."],
    ["Time-series acceleration", "Timescale hypertables and continuous aggregate views", "1 minute, 15 minute, 1 hour, and 1 day views support KPI/report/dashboard queries."],
    ["Analytics metadata", "energy_baselines, anomalies, energy_tariffs, carbon_factors, audit_log", "Model, anomaly, cost, carbon, and audit support tables."],
    ["ISO 50001 structures", "energy_sources, seus, seu_energy_performance, enpi_baselines, enpi_performance, energy_targets, action_plans", "Energy-source, significant-energy-use, EnPI, target, and action-plan concepts."],
    ["Model/forecast tracking", "model_performance tables, energy_forecasts, model_training_history, model_alerts", "Forecast output and model lifecycle/performance observations."],
]


CONFIG_GROUP_ROWS = [
    ["Database", "POSTGRES_USER, POSTGRES_PASSWORD, POSTGRES_DB, POSTGRES_HOST, POSTGRES_PORT", ".env.example; docker-compose.yml; analytics/config.py"],
    ["Service ports", "NGINX_HTTP_PORT, ANALYTICS_PORT, SIMULATOR_PORT, GRAFANA_PORT, NODERED_PORT, RASA_PORT, CHATBOT_PORT, REDIS_EXTERNAL_PORT", ".env.example; docker-compose.yml"],
    ["Security and auth", "JWT_SECRET, API_KEY, ADMIN_EMAILS, JWT expiration settings, Node-RED password hash/credential secret", ".env.example; auth-service/auth_service.py; nodered/settings.js"],
    ["Telemetry", "MQTT_HOST, MQTT_PORT, MQTT_USERNAME, MQTT_PASSWORD, simulator intervals and anomaly settings", ".env.example; simulator/config.py"],
    ["Analytics behavior", "LOG_LEVEL, scheduler settings, model storage, anomaly thresholds, cost/carbon defaults", "analytics/config.py; .env.example"],
    ["OVOS proxy/runtime", "OVOS_BRIDGE_HOST, OVOS_BRIDGE_PORT, OVOS_BRIDGE_TIMEOUT, OVOS external ports", ".env.example; analytics/api/routes/ovos_voice.py; OVOS-EnMS repository compose"],
]


ERROR_HANDLING_ROWS = [
    ["Analytics request handling", "Request logging middleware, validation exception handling, timeout middleware, and generic exception handling are present in analytics/main.py and middleware modules."],
    ["Analytics route behavior", "Routes generally return structured JSON errors or raise HTTPException when validation, lookup, or database work fails."],
    ["Auth service", "Endpoints catch exceptions, log errors, and return success/error JSON while auth_service.py performs bcrypt/JWT verification and audit logging."],
    ["Simulator", "Lifecycle startup/shutdown, MQTT/database connection handling, and route-level start/stop/status errors are implemented."],
    ["OVOS bridge/client", "Bridge returns structured failure responses on messagebus/query errors; ENMSClient avoids retrying ordinary 4xx errors and retries transient/timeouts/5xx."],
]


KPI_CATALOG_ROWS = [
    ["SEC", "Supported by implementation", "calculate_sec SQL function and /api/v1/kpi/sec route", "Requires energy and production aggregate data for the selected period."],
    ["Peak demand", "Supported by implementation", "calculate_peak_demand SQL function and route", "Uses 15-minute aggregate peak_demand_kw."],
    ["Load factor", "Supported by implementation", "calculate_load_factor SQL function and route", "Depends on average/max power availability."],
    ["Energy cost", "Supported by implementation", "calculate_energy_cost SQL function and service wrapper", "Uses active tariff rows when present and default fallback in SQL."],
    ["Carbon intensity/emissions", "Supported by implementation", "calculate_carbon_intensity SQL function and carbon route", "Uses carbon_factors with fallback factor; not independent emissions assurance."],
    ["Factory KPI rollups", "Implemented/configured", "/api/v1/kpi/factory/{factory_id} and /api/v1/kpi/factories", "Some route-level estimates use constants or aggregate assumptions; formal audit use requires independent validation."],
    ["Model performance KPIs", "Implemented/configured", "model_performance routes and dashboards", "R2/RMSE/MAPE-style metrics depend on recorded model history."],
    ["Operational efficiency/OEE", "Configured dashboard/reporting view", "Grafana operational-efficiency dashboards and production route", "Formal operational reporting requires validation of the dashboard SQL, source data, and reporting definitions."],
]


REPORT_WORKFLOW_ROWS = [
    ["Legacy report request", "Client calls /api/v1/reports/types, /preview, or /generate for monthly_enpi.", "analytics/api/routes/reports.py"],
    ["Legacy data assembly", "MonthlyEnPIReport builds summary, machine metrics, EnPI values, targets, achievements, and charts.", "analytics/reports/monthly_enpi_report.py"],
    ["Legacy output", "ReportLab PDF is returned by the route.", "analytics/reports/base_report.py; analytics/api/routes/reports.py"],
    ["V2 report request", "Client calls /api/v1/reports/v2/generate and later /v2/download/{report_id}.", "analytics/api/routes/reports.py"],
    ["V2 generation", "ReportService coordinates data fetch, components, charts, HTML/PDF generation, and temporary output path.", "analytics/reports_v2/services/report_service.py"],
    ["V2 scope boundary", "Some data fetcher/service values are estimated or placeholder-like.", "analytics/reports_v2/services/data_fetcher.py; report_service.py"],
]


DATA_QUALITY_ROWS = [
    ["Clock/time range", "KPI, baseline, forecast, and dashboard results assume timestamps are correctly generated and synchronized."],
    ["Telemetry completeness", "SEC and production-linked KPIs require both energy and production data; missing production affects denominator quality."],
    ["Topic consistency", "Node-RED routing assumes MQTT topics match expected factory/# structure and payload type handling."],
    ["Tariff/factor validity", "Cost and carbon outputs depend on active tariff and carbon-factor records or configured fallback factors."],
    ["Simulator vs real data", "Seed/demo simulator data is suitable for demonstration but should be separated from live factory evidence."],
    ["Aggregate freshness", "Continuous aggregates and dashboards depend on database refresh behavior and current ingested data."],
]


AUDIT_CAUTION_ROWS = [
    ["Audit-grade KPI use", "Generated reports are operational outputs. Formal audit use requires validation of formulas, data sources, tariff/factor records, and reporting period boundaries."],
    ["Carbon reporting", "Carbon/emissions values are implemented where functions/routes/dashboards exist, but formal emissions reporting requires verified factors, scope definitions, and governance outside this codebase."],
    ["Demo seed data", "Seeded factories and machines should be labeled as demonstration data unless replaced by real facility data."],
    ["Estimated calculations", "Where V2 reports or dashboard panels derive estimates, classify them as operational estimates rather than certified calculations."],
]


HEALTHCHECK_ROWS = [
    ["nginx", "wget/curl style check to /health", "Confirms gateway process responds, not full upstream health."],
    ["postgres", "pg_isready", "Confirms database accepts connections."],
    ["mqtt", "mosquitto_pub/sub or broker health command", "Confirms broker is reachable with configured credentials."],
    ["redis", "redis-cli ping with password", "Confirms Redis responds."],
    ["simulator", "HTTP /health", "Confirms simulator API process health."],
    ["nodered", "HTTP admin/API health check", "Confirms Node-RED runtime responds."],
    ["grafana", "HTTP /api/health", "Confirms Grafana process health."],
    ["analytics", "HTTP /api/v1/health", "Confirms analytics service and database checks exposed by route."],
    ["auth-service", "HTTP /api/auth/health", "Confirms Flask auth service responds."],
    ["rasa/rasa-actions/chatbot", "HTTP health or root checks from Compose healthcheck", "Confirms chatbot components respond."],
]


BACKUP_RECOVERY_ROWS = [
    ["PostgreSQL/TimescaleDB", "Use pg_dump/pg_restore or platform database backups before upgrades and before docker compose down -v.", "No tracked generic database backup script exists in production."],
    ["Grafana dashboards", "Keep intended dashboard JSON under grafana/dashboards and provisioning under grafana/provisioning.", "Back up runtime edits before replacing Grafana volumes."],
    ["Node-RED flows", "nodered/data/flows.json is tracked; credential-bearing runtime files should not be published.", "Credential files are excluded from documentation."],
    ["Docker volumes", "postgres-data, grafana-data, mqtt data/logs, redis-data, Node-RED data", "Volume deletion is destructive."],
    ["Documentation package", "Regenerate DOCX from docs/final-delivery/source/generate_delivery_docs.py and source Markdown/assets.", "Keeps source and deliverables reproducible."],
]


UPGRADE_REDEPLOY_ROWS = [
    ["Pre-upgrade", "Confirm current commit, back up database/volumes, export dashboard changes, record .env keys without exposing values."],
    ["Build", "docker compose build after pulling source changes or switching approved delivery bundles."],
    ["Redeploy", "docker compose up -d; use --wait when supported by Docker Compose."],
    ["Smoke test", "Run docker compose ps, service logs, /health endpoints, analytics /api/v1/health, and verify.sh when services are running."],
    ["Rollback", "Return to the previous commit or approved delivery version and restore volumes/database backup if schema/data changed."],
]


HARDENING_ROWS = [
    ["Credentials", "Rotate generated setup secrets; use strong admin, database, Redis, MQTT, JWT, API, and Node-RED values."],
    ["Network exposure", "Restrict direct service ports; expose only intended Nginx/TLS routes to external users."],
    ["TLS", "Configure certificates in Nginx or an upstream reverse proxy before internet-facing use."],
    ["Runtime users", "Review container users and file permissions; OVOS Dockerfile creates a non-root ovos user."],
    ["Backups", "Implement tested backup/restore for database, dashboard changes, and operational volumes."],
    ["Monitoring", "Add external monitoring/log collection for health endpoints, disk usage, restart counts, and data freshness."],
    ["Secrets hygiene", "Never publish .env, runtime credential files, tokens, database dumps, or logs containing credentials."],
]


TROUBLESHOOTING_COMMAND_ROWS = [
    ["List service state", "docker compose ps", "Shows container state, health, and exposed ports."],
    ["Inspect logs", "docker compose logs <service> --tail=100", "Use owning service first: nginx, analytics, postgres, mqtt, nodered, grafana, auth-service, chatbot, rasa."],
    ["Validate configuration", "docker compose config --quiet", "Catches Compose syntax/interpolation errors."],
    ["Gateway health", "curl -fsS http://localhost:8080/health", "Expected healthy response from Nginx when stack is running."],
    ["Analytics health", "curl -fsS http://localhost:8001/api/v1/health", "Checks analytics service directly."],
    ["Verification script", "HUMANERDIA_BASE_URL=... ANALYTICS_BASE_URL=... OVOS_BASE_URL=... ./verify.sh", "Runs live checks; skips OVOS if bridge is unreachable."],
]


SKILL_BACKEND_MAPPING_ROWS = [
    ["Factory overview", "factory_summary", "/factory/summary", "Factory-wide status, energy, cost, machine, and anomaly summary."],
    ["Machine status", "get_machine_status", "/machines/status/{machine_name}", "Current machine state and related statistics."],
    ["Machine list", "list_machines", "/machines", "Available/active machine inventory and machine-name discovery."],
    ["Energy query", "get_energy_timeseries, get_latest_reading", "/timeseries/energy, /timeseries/latest/{machine_id}", "Historical or latest energy data after machine lookup."],
    ["Power query", "get_power_timeseries, get_machine_status", "/timeseries/power, /machines/status/{machine_name}", "Current or historical power answer."],
    ["Ranking/top consumers", "get_top_consumers", "/analytics/top-consumers or /ovos/top-consumers", "Top/bottom consumers by supported metric."],
    ["Anomalies", "get_recent_anomalies, detect_anomalies", "/anomaly/recent, /anomaly/detect", "Recent or active anomaly information."],
    ["KPIs", "get_all_kpis and KPI-specific methods", "/kpi/all and KPI routes", "SEC, peak demand, load factor, cost, carbon and related rollups."],
    ["Baselines/drivers", "predict_baseline, list_baseline_models, get_baseline_drivers", "/baseline/*", "Baseline prediction, models, and energy-driver explanations."],
    ["Reports", "get_report_types, preview_report, generate_report", "/reports/types, /reports/preview, /reports/generate", "Report discovery, preview, and generation."],
]


SKILL_CONFIG_REFERENCE_ROWS = [
    ["ENMS_API_URL", "Container environment variable", "Backend API base URL for the OVOS runtime/bridge environment."],
    ["OVOS_BRIDGE_PORT", "Container environment variable", "REST bridge listen port; default 5000."],
    ["OVOS_MESSAGEBUS_PORT", "Compose port", "Messagebus port exposed by OVOS-EnMS Compose; default 8181."],
    ["STRUCTURED_RESPONSE_GRACE_SECONDS", "Bridge environment variable", "Additional wait for structured enms.skill.response payload after speech event."],
    ["OVOS_TTS_ENABLED", "Bridge/runtime environment variable", "Controls TTS behavior in the OVOS runtime."],
    ["llm_model_path", "Skill setting", "Configured GGUF path for optional LLM parser fallback."],
    ["confidence_threshold", "Skill/validator setting", "Minimum confidence for accepted parsed intents; default observed value is 0.85."],
    ["enable_fuzzy_matching", "Skill/validator setting", "Allows fuzzy machine-name matching and suggestions."],
    ["api_timeout_seconds / api_max_retries", "Skill settings", "Backend request timeout and retry behavior."],
]


SKILL_FAILURE_ROWS = [
    ["Messagebus unavailable", "REST bridge health reports disconnected state; query handling cannot complete normal OVOS event round trip."],
    ["Backend API unavailable", "ENMSClient logs request/connect errors; skill should return failure/clarification rather than fabricated KPI values."],
    ["Low confidence parse", "Validator rejects output below threshold and suggests rephrasing."],
    ["Unknown machine", "Validator rejects invalid names and can suggest fuzzy matches."],
    ["Ambiguous comparison", "Validator expands groups when possible or asks for clarification when insufficient machines match."],
    ["Template failure", "Skill has fallback response generation for several intent/data shapes."],
    ["LLM dependencies/model missing", "Hybrid parser continues with heuristic/Adapt tiers and clarification fallback; LLM fallback is optional."],
]


SKILL_EXAMPLE_ROWS = [
    ["Energy", "How much energy did Compressor-1 use yesterday?"],
    ["Power", "What is the current power of Boiler-1?"],
    ["Status", "Is HVAC-Main running?"],
    ["Overview", "Give me a factory overview."],
    ["Ranking", "Show the top three energy consumers today."],
    ["Comparison", "Compare Compressor-1 and Compressor-EU-1."],
    ["Anomalies", "Any active anomalies today?"],
    ["Forecast", "What is tomorrow's demand forecast?"],
    ["Baseline", "What is the baseline for Injection-Molding-1?"],
    ["Drivers", "Explain the energy drivers for Compressor-1."],
    ["SEUs", "List significant energy uses."],
    ["Reports", "Generate a monthly energy report."],
    ["Health/help", "What can you do? Is the system healthy?"],
]


FINAL_ARTIFACT_ROWS = [
    ["Final DOCX reports", "Six stakeholder deliverables in docs/final-delivery/."],
    ["Source Markdown", "docs/final-delivery/source/*.md mirrors the generated DOCX content for maintainability."],
    ["Evidence map", "docs/final-delivery/source/evidence-map.md maps recurring claims to source files and validation output."],
    ["Generation script", "docs/final-delivery/source/generate_delivery_docs.py regenerates DOCX, Markdown, and diagrams."],
    ["Diagram assets", "docs/final-delivery/assets/*.png contains architecture, data-flow, deployment, topology, and OVOS lifecycle diagrams."],
    ["Application source", "HumanEnerDIA production source tree plus separate OVOS-EnMS repository remain the authoritative technical sources."],
]


ACCEPTANCE_CHECK_ROWS = [
    ["Documentation files exist", "All six required DOCX files are present under docs/final-delivery/."],
    ["Traceability", "Source references point to tracked delivery files, implementation code, configuration, and validation evidence."],
    ["Compose validation", "docker compose config --quiet passes for HumanEnerDIA production; OVOS-EnMS Compose validates separately."],
    ["DOCX integrity", "All DOCX files open as valid ZIP/DOCX packages and contain embedded media where expected."],
    ["Secrets hygiene", "Generated docs and sources are scanned for sensitive placeholders and private values."],
    ["Scope boundaries", "Runtime validation, audit use, OVOS deployment, and production hardening are clearly defined."],
]


DEMO_READINESS_ROWS = [
    ["Before demo", "Run docker compose ps, verify.sh, Nginx/analytics health checks, and OVOS /health if assistant demo is planned."],
    ["Data freshness", "Confirm simulator or real ingestion is producing recent records before showing dashboards/KPIs."],
    ["Credentials", "Use prepared demo/operator credentials without displaying .env or secrets."],
    ["Dashboards", "Open Grafana dashboards and confirm panels load with current data."],
    ["Assistant", "Run at least one machine status query and one KPI/report-style query through the OVOS bridge if OVOS is in scope."],
    ["Known cautions", "Be ready to explain demo data, partial V2 report semantics, and production hardening steps."],
]


def load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> List[str]:
    words = text.split()
    lines: List[str] = []
    current: List[str] = []
    for word in words:
        trial = " ".join(current + [word])
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_width or not current:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines


def draw_box(draw: ImageDraw.ImageDraw, box: Tuple[int, int, int, int], title: str, body: str, fill: Tuple[int, int, int], outline: Tuple[int, int, int]) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=3)
    title_font = load_font(24, True)
    body_font = load_font(17)
    max_width = x2 - x1 - 36

    y = y1 + 14
    for line in wrap_text(draw, title, title_font, max_width):
        draw.text((x1 + 18, y), line, fill=(20, 32, 45), font=title_font)
        y += 30

    y += 8
    for line in wrap_text(draw, body, body_font, max_width):
        if y + 19 > y2 - 12:
            break
        draw.text((x1 + 18, y), line, fill=(42, 54, 68), font=body_font)
        y += 23


def draw_arrow(draw: ImageDraw.ImageDraw, start: Tuple[int, int], end: Tuple[int, int], color=(71, 85, 105)) -> None:
    draw.line([start, end], fill=color, width=4)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) > abs(ey - sy):
        sign = 1 if ex > sx else -1
        points = [(ex, ey), (ex - sign * 14, ey - 9), (ex - sign * 14, ey + 9)]
    else:
        sign = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - sign * 14), (ex + 9, ey - sign * 14)]
    draw.polygon(points, fill=color)


def make_diagram(filename: str, title: str, boxes: Sequence[Tuple[Tuple[int, int, int, int], str, str]], arrows: Sequence[Tuple[Tuple[int, int], Tuple[int, int]]], size=(1500, 900)) -> str:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", size, (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_font = load_font(38, True)
    draw.text((50, 32), title, fill=(15, 23, 42), font=title_font)
    for start, end in arrows:
        draw_arrow(draw, start, end)
    palette = [
        ((226, 232, 240), (71, 85, 105)),
        ((219, 234, 254), (37, 99, 235)),
        ((220, 252, 231), (22, 101, 52)),
        ((254, 243, 199), (180, 83, 9)),
        ((243, 232, 255), (126, 34, 206)),
    ]
    for idx, (box, title_text, body) in enumerate(boxes):
        fill, outline = palette[idx % len(palette)]
        draw_box(draw, box, title_text, body, fill, outline)
    path = ASSET_DIR / filename
    img.save(path)
    return str(path)


def generate_diagrams() -> dict:
    diagrams = {}
    diagrams["system_context"] = make_diagram(
        "system-context.png",
        "System Context",
        [
            ((70, 150, 360, 310), "Users and Operators", "Portal, Grafana, analytics UI, Rasa chatbot, and OVOS voice/text queries"),
            ((470, 130, 780, 330), "Nginx Gateway", "Public HTTP gateway for portal, APIs, Grafana, Node-RED, auth, chatbot, and OVOS proxy paths"),
            ((880, 100, 1240, 250), "HumanEnerDIA Core", "Analytics, simulator, ingestion, auth, chatbot/Rasa, dashboards"),
            ((880, 330, 1240, 500), "Data and Messaging", "TimescaleDB/PostgreSQL, MQTT, Redis, provisioning, persistent volumes"),
            ((470, 520, 780, 710), "OVOS-EnMS Boundary", "Separate assistant runtime: bridge, messagebus, skill, parser, validator, API client"),
        ],
        [
            ((360, 230), (470, 230)),
            ((780, 230), (880, 180)),
            ((1060, 250), (1060, 330)),
            ((630, 520), (880, 440)),
            ((780, 615), (900, 240)),
        ],
    )
    diagrams["telemetry_flow"] = make_diagram(
        "telemetry-data-flow.png",
        "Telemetry and Analytics Data Flow",
        [
            ((70, 150, 330, 300), "Simulator / Devices", "Publishes factory telemetry streams to MQTT"),
            ((430, 150, 690, 300), "MQTT Broker", "Mosquitto receives factory/# telemetry streams"),
            ((790, 150, 1060, 330), "Node-RED", "Parses topics, validates payloads, and writes normalized records"),
            ((1160, 150, 1440, 330), "TimescaleDB", "Raw hypertables plus 1min, 15min, 1hour, and 1day aggregates"),
            ((790, 500, 1060, 680), "Analytics APIs", "KPIs, baselines, forecasts, anomalies, reports, ISO 50001"),
            ((1150, 500, 1450, 690), "Dashboards and Reports", "Grafana, analytics UI, reports, Rasa/OVOS consumers"),
        ],
        [
            ((330, 225), (430, 225)),
            ((690, 225), (790, 225)),
            ((1060, 240), (1160, 240)),
            ((1300, 330), (1060, 540)),
            ((1060, 590), (1160, 590)),
        ],
    )
    diagrams["ovos_lifecycle"] = make_diagram(
        "ovos-query-lifecycle.png",
        "OVOS Query Lifecycle",
        [
            ((50, 150, 290, 310), "User / Portal", "Natural-language text or voice-derived request"),
            ((340, 150, 580, 330), "REST Bridge", "POST /query or /query/voice emits recognizer_loop:utterance"),
            ((630, 150, 870, 330), "OVOS Messagebus", "Carries utterance, speak event, and enms.skill.response payloads"),
            ((920, 150, 1160, 330), "EnMS Skill", "Parser, validator, API client, handlers, context, formatter"),
            ((1210, 150, 1450, 330), "HumanEnerDIA API", "Configured HumanEnerDIA-compatible /api/v1 backend"),
            ((1210, 500, 1450, 680), "Response", "Bridge returns voice text, data, insights, and optional report metadata"),
        ],
        [
            ((290, 230), (340, 230)),
            ((580, 230), (630, 230)),
            ((870, 230), (920, 230)),
            ((1160, 230), (1210, 230)),
            ((1330, 330), (1330, 500)),
        ],
    )
    diagrams["deployment_flow"] = make_diagram(
        "deployment-startup-flow.png",
        "Deployment Startup Flow",
        [
            ((70, 150, 360, 310), "Prepare Host", "Docker Engine and Compose v2, production repository, .env.example available"),
            ((470, 150, 760, 330), "setup.sh", "Creates .env if needed, generates first-run secrets, updates URLs, detects optional OVOS compose file"),
            ((870, 150, 1160, 330), "Compose Validation", "docker compose config validates base stack and optional compose files"),
            ((470, 510, 760, 690), "Build and Start", "docker compose build and up -d, optionally with --wait"),
            ((870, 510, 1160, 690), "Verify", "Health endpoints, verify.sh, smoke query when OVOS is deployed"),
        ],
        [
            ((360, 230), (470, 230)),
            ((760, 240), (870, 240)),
            ((1010, 330), (760, 560)),
            ((760, 600), (870, 600)),
        ],
    )
    diagrams["docker_topology"] = make_diagram(
        "docker-service-topology.png",
        "Docker Service Topology",
        [
            ((60, 150, 330, 310), "Host Access", "Browser/API users reach Nginx on 8080. Direct ports are mapped for selected operations services"),
            ((420, 130, 700, 330), "Nginx", "Routes portal, analytics, Grafana, Node-RED, auth, chatbot, simulator, OVOS proxy paths"),
            ((790, 130, 1090, 330), "App Services", "analytics, auth-service, chatbot, Rasa, rasa-actions, simulator"),
            ((1180, 130, 1460, 330), "Data Services", "PostgreSQL/TimescaleDB, MQTT broker, Redis, named persistent volumes"),
            ((420, 500, 700, 680), "Operator UIs", "Grafana dashboards, Node-RED editor, analytics UI, static portal"),
            ((790, 500, 1090, 680), "OVOS-EnMS Boundary", "Separate assistant runtime; production base Compose has no OVOS service"),
        ],
        [
            ((330, 230), (420, 230)),
            ((700, 230), (790, 230)),
            ((1090, 230), (1180, 230)),
            ((560, 330), (560, 500)),
            ((940, 500), (940, 330)),
            ((700, 590), (790, 590)),
        ],
    )
    return diagrams


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def keep_cell_with_next(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.keep_with_next = True


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        shade_cell(header_cells[idx], "D9EAF7")
    set_repeat_table_header(table.rows[0])
    set_cant_split_row(table.rows[0])
    data_rows = []
    for row in rows:
        new_row = table.add_row()
        data_rows.append(new_row)
        set_cant_split_row(new_row)
        cells = new_row.cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if len(data_rows) >= 2:
        for cell in data_rows[-2].cells:
            keep_cell_with_next(cell)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(26)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(30, 64, 175)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def add_footer(doc: Document, title: str) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"{PROJECT_NAME} | {title} | Version {DOC_VERSION} | {DOC_DATE}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 116, 139)


def add_title_page(doc: Document, spec: DocSpec) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT_NAME)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 64, 175)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(spec.title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    rows = [
        ["Version", DOC_VERSION],
        ["Date", DOC_DATE],
        ["Status", DOC_STATUS],
        ["Purpose", spec.purpose],
        ["Intended audience", spec.audience],
    ]
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            set_cell_text(meta.rows[row_idx].cells[col_idx], value, bold=(col_idx == 0))
            if col_idx == 0:
                shade_cell(meta.rows[row_idx].cells[col_idx], "E2E8F0")
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Source basis: ")
    r.bold = True
    note.add_run(spec.evidence_note)
    doc.add_page_break()


def add_toc(doc: Document, spec: DocSpec) -> None:
    doc.add_heading("Table of Contents", level=1)
    for idx, section in enumerate(spec.sections, start=1):
        doc.add_paragraph(f"{idx}. {section.title}")
        for sub_idx, sub in enumerate(section.subsections, start=1):
            doc.add_paragraph(f"{idx}.{sub_idx} {sub.title}", style="List Bullet")
    doc.add_page_break()


def add_text_block(paragraph, block: TextBlock) -> None:
    if isinstance(block, tuple):
        label, text = block
        run = paragraph.add_run(label.rstrip())
        run.bold = True
        paragraph.add_run(f" {text.lstrip()}")
    else:
        paragraph.add_run(block)


def add_section(doc: Document, section: Section, level: int = 1) -> None:
    doc.add_heading(section.title, level=level)
    if section.figure:
        image_name, caption = section.figure
        image_path = ASSET_DIR / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(7.2))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(8.5)
    for para in section.paragraphs:
        p = doc.add_paragraph()
        add_text_block(p, para)
    for bullet in section.bullets:
        doc.add_paragraph(bullet, style="List Bullet")
    if section.table:
        intro = table_intro_for(section.title)
        if intro:
            p = doc.add_paragraph()
            add_text_block(p, intro)
        add_table(doc, section.table[0], section.table[1])
    followup_blocks = section.table_followup or ([table_followup_for(section.title)] if section.table else [])
    for para in followup_blocks:
        p = doc.add_paragraph()
        add_text_block(p, para)
    for sub in section.subsections:
        add_section(doc, sub, level=min(level + 1, 3))


def build_document(spec: DocSpec) -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc, spec)
    add_toc(doc, spec)
    for section in spec.sections:
        add_section(doc, section)
    add_footer(doc, spec.title)
    output = OUT_DIR / spec.filename
    doc.save(output)


TABLE_INTROS = {
    "System Context": "The table below defines the delivery boundaries so readers can distinguish the EnMS platform, the companion OVOS runtime, and the external users or integrators that interact with them.",
    "Component Responsibilities": "Read this table as a responsibility map: each row identifies the architectural owner for a capability and the source material that supports that assignment.",
    "Runtime Services": "The service inventory explains what runs in the base Compose deployment and what role each service plays in the operational stack.",
    "Dependency And Request Flow": "The dependency flows show how requests and data move through the stack from the point of entry to storage, analytics, dashboards, and assistant responses.",
    "Data And Message Flow": "This table follows telemetry through each processing stage, from publication to storage and then to downstream consumers.",
    "External And Internal Interfaces": "The interface list identifies the main access routes and clarifies which paths are intended for browsers, APIs, operations tools, or the optional assistant bridge.",
    "Deployment Variants And Boundaries": "The deployment variants table separates evaluation, base production, companion OVOS, and hardened production responsibilities.",
    "Operational Risks": "The operational risk table summarizes deployment choices that require operator governance in a production environment.",
    "Limitations And Assumptions": "The limitations table defines scope boundaries that affect validation, audit use, optional assistant behavior, and production readiness.",
    "Source References": "The source reference table links the document's major claims to the tracked files or validation evidence used to support them.",
    "Module Responsibility Matrix": "This matrix maps software areas to implementation responsibilities so maintainers can quickly identify the owning subsystem for each capability.",
    "Analytics Service Design": "The endpoint table groups the analytics API by domain capability and points to the router modules that implement each area.",
    "API Design Details": "The API design table explains how routes, models, database access, compatibility paths, and UI-serving routes are structured.",
    "Database And Schema Design": "This table summarizes the first-start schema groups and shows how each database area supports the EnMS domain model.",
    "Database Object Catalog": "The object catalog gives a stakeholder-level view of database groups without reproducing every column definition.",
    "Service-Layer Design": "The service-layer table highlights where calculation, modeling, reporting, and event behavior live beyond the route handlers.",
    "Simulator And Ingestion Design": "This table connects simulator behavior, MQTT publication, and Node-RED ingestion into one implementation view.",
    "Authentication And Authorization Design": "The authentication table identifies the implemented account, session, verification, admin, and audit capabilities.",
    "Configuration Reference": "The configuration table groups environment-driven settings without exposing private runtime values.",
    "Error Handling And Logging": "The error-handling table describes how the main services surface failures and logs at runtime.",
    "Known Design Gaps And Placeholders": "This table states known design gaps in external delivery language so the current scope is clear without overstating completeness.",
    "Deployment And Configuration": "The table below summarizes the runtime configuration points that determine how OVOS connects to the EnMS backend.",
    "Supported Intent And Query Families": "The intent family table explains the categories of operational questions represented in the skill code.",
    "Intent Parsing And Routing": "The parser table shows the ordered routing strategy from fast deterministic handling to optional fallback parsing.",
    "Backend API Client And Adapter Behavior": "The backend-client table shows the EnMS API areas the skill can call after parsing and validation.",
    "Backend Method Mapping": "This mapping connects natural-language categories to backend methods and API areas for implementation traceability.",
    "Example Supported Queries": "The example table provides representative phrases that illustrate supported query families and expected usage patterns.",
    "Failure Behavior": "The failure table explains how the assistant should behave when parsing, validation, backend connectivity, or optional components are unavailable.",
    "Energy Data Model": "The data model table identifies the key records used for energy management, reporting, analytics, and ISO 50001-oriented workflows.",
    "KPI Formula Evidence": "The formula table separates implemented KPI calculations from higher-level dashboard or report presentation.",
    "KPI Catalog And Classification": "The classification table helps readers distinguish calculation-backed KPIs, configured views, and measures that require governance for formal use.",
    "Analytics Endpoints And Modules": "This table groups the reporting and analytics routes by capability area so readers can connect dashboard/report features to API implementation.",
    "Grafana Dashboard Capabilities": "The dashboard table summarizes the configured Grafana views and the operational themes each view is designed to support.",
    "Node-RED Ingestion Pipeline": "The ingestion table explains how MQTT payloads are routed, validated, transformed, and written into PostgreSQL.",
    "Report Generation Capabilities": "The report table distinguishes the legacy EnPI report path from the newer V2 PDF workflow.",
    "Report Workflow": "The workflow table follows a report request from API entry point through data assembly, rendering, and output.",
    "Data-Quality Assumptions": "This table identifies the assumptions that determine whether KPI, dashboard, and report outputs are meaningful.",
    "Audit-Use Cautions": "The audit-use table defines the additional validation needed before using outputs for formal regulatory, financial, or assurance purposes.",
    "Implemented, Configured, Partial, And Demo Data Distinctions": "The distinction table classifies capabilities by evidence level so readers can separate implemented behavior from configured views and demo data.",
    "Deployment Prerequisites": "The prerequisite table defines the host and tooling conditions needed before installation or startup.",
    "Docker Compose Services": "The service table is the operator's runtime inventory for base HumanEnerDIA deployment.",
    "Networks, Volumes, And Ports": "This table identifies the persistent and network resources that influence backup, firewall, and redeployment planning.",
    "Environment Variable Groups": "The environment table groups configuration by operational concern while avoiding private runtime values.",
    "Step-By-Step Deployment": "The deployment table gives a controlled sequence for preparing, validating, building, starting, and verifying the stack.",
    "Startup, Shutdown, And Clean Reinstall Procedures": "This procedure table separates routine operations from destructive reinstall actions.",
    "Verification Scripts And Health Checks": "The verification table clarifies which checks are static validation and which require live services.",
    "Healthcheck Details": "The healthcheck table explains what each container health signal proves and what it does not prove.",
    "Backup And Recovery": "The backup table identifies the assets that must be protected before upgrades, destructive cleanup, or production use.",
    "Upgrade And Redeployment": "The redeployment table describes a controlled operator sequence for updating the system.",
    "Troubleshooting Commands": "The troubleshooting command table lists safe inspection commands for diagnosing a running deployment.",
    "Production Hardening Checklist": "The hardening table converts deployment risk areas into operator actions for production readiness.",
    "Troubleshooting Scenarios": "The scenario table connects common symptoms to the first services and checks operators should inspect.",
    "Handover Summary": "The handover table explains how the six documents work together as a final delivery set.",
    "Delivery Artifact List": "The artifact table identifies the maintained deliverables and the source materials that keep them reproducible.",
    "Quick-Start Summary": "The quick-start table condenses installation and verification into a high-level sequence for orientation.",
    "Installation And Access": "The access table shows the default endpoints readers can use after deployment and verification.",
    "Main Workflows": "The workflow table summarizes the operational paths users and operators will exercise most often.",
    "Analytics, Dashboards, Reports, And Assistants": "This table distinguishes dashboards, analytics UI, text chatbot behavior, and OVOS operational assistant behavior.",
    "Maintenance And Troubleshooting": "The maintenance table summarizes recurring operator responsibilities after installation.",
    "Acceptance Checklist": "The acceptance table records the deliverable and traceability checks used to confirm the documentation set.",
    "Demo-Readiness Checklist": "The demo-readiness table lists live checks that must be completed on the actual target deployment before a demonstration.",
}


TABLE_FOLLOWUPS = {
    "System Context": "Operationally, this boundary view prevents the EnMS platform, text chatbot, and OVOS assistant from being described as one inseparable runtime. That distinction matters when deploying only EnMS, only OVOS, or the full stack.",
    "Component Responsibilities": "The responsibilities show that HumanEnerDIA is modular: gateway, storage, ingestion, analytics, visualization, authentication, chatbot, and assistant functions can be discussed and verified separately.",
    "Runtime Services": "This inventory confirms that the base EnMS deployment is a multi-service stack. Operators should use it to plan resource allocation, health checks, firewall policy, and troubleshooting ownership.",
    "Dependency And Request Flow": "The flow relationships show where failures will propagate. For example, stale telemetry affects dashboards and assistant answers, while gateway routing issues can affect multiple browser-facing paths.",
    "Data And Message Flow": "The data flow emphasizes that KPIs and reports depend on successful ingestion, storage, and aggregate refresh. It also clarifies where external device integration would connect.",
    "External And Internal Interfaces": "The interface list supports access planning. Public deployments should prefer controlled gateway routes and restrict direct service ports unless an operational need is approved.",
    "Deployment Variants And Boundaries": "These variants allow WASABI stakeholders and operators to choose the correct package without assuming that OVOS is always bundled into the EnMS runtime.",
    "Operational Risks": "The listed risks are manageable with normal production controls: firewalling, credential rotation, backups, monitoring, and explicit assistant deployment.",
    "Source References": "These references provide traceability for technical claims. They are intended to support maintenance and verification without exposing secrets or local runtime state.",
    "Module Responsibility Matrix": "The matrix also gives maintainers a change-impact map: API changes, schema changes, ingestion changes, and authentication changes have different owners and test surfaces.",
    "Analytics Service Design": "The route grouping shows that analytics is the main domain API. Changes to these routes can affect dashboards, reports, portal views, and assistant answers.",
    "API Design Details": "The design details show a pragmatic API surface: typed routes and service-backed routes coexist with compatibility paths and UI-serving endpoints.",
    "Database And Schema Design": "The schema design supports both operational telemetry and higher-level energy management concepts. It is the foundation for dashboards, KPIs, reports, and assistant responses.",
    "Database Object Catalog": "This catalog is useful for handover because it shows where core records, time-series facts, ISO 50001 concepts, and model tracking are stored.",
    "Service-Layer Design": "The service layer reduces duplication and keeps complex behavior closer to domain modules, while some legacy routes still contain direct SQL or local calculations.",
    "Simulator And Ingestion Design": "This design enables repeatable demonstrations and also shows the expected integration pattern for real telemetry producers.",
    "Authentication And Authorization Design": "The authentication design provides practical account and admin controls, while enterprise identity integration remains a target-environment responsibility.",
    "Configuration Reference": "These groups help operators plan configuration review without exposing the deployed values themselves.",
    "Error Handling And Logging": "The error-handling view supports operational diagnosis. It also shows where callers can expect structured errors versus log-based investigation.",
    "Known Design Gaps And Placeholders": "These gaps define current scope boundaries. They should guide future hardening and validation work rather than obscure the implemented capabilities.",
    "Deployment And Configuration": "Correct backend URL configuration is central to OVOS readiness: the assistant can answer operational questions only when it can reach the EnMS-compatible API.",
    "Configuration Reference": "Together, these settings define how the assistant connects, listens, validates, and optionally uses fallback parsing. Runtime values should be controlled by the deployment owner.",
    "Supported Intent And Query Families": "The intent families show broad operational coverage while still depending on parser confidence, backend data, and machine names present in the target system.",
    "Intent Parsing And Routing": "This layered routing model keeps common questions fast and deterministic, while optional fallback parsing expands coverage when properly configured.",
    "Backend API Client And Adapter Behavior": "The client layer is the bridge between language understanding and EnMS data. Its retry behavior helps with transient backend issues without hiding persistent configuration errors.",
    "Backend Method Mapping": "The mapping gives integrators a practical checklist for adapter compatibility when connecting OVOS to another EnMS backend.",
    "Example Supported Queries": "The examples should be used as smoke-test and training phrases, with expected answers determined by the data currently available in the EnMS backend.",
    "Failure Behavior": "The failure behavior supports trustworthy operation: unclear or unsupported inputs should lead to clarification or explicit failure rather than fabricated operational values.",
    "Energy Data Model": "This model shows that useful energy reporting depends on aligned factory, machine, telemetry, tariff, carbon, and production records.",
    "KPI Formula Evidence": "The formulas provide traceable calculation evidence for key metrics. They also identify which metrics depend on production counts, tariff records, carbon factors, or aggregate freshness.",
    "KPI Catalog And Classification": "The classification helps stakeholders interpret dashboards and reports with the appropriate confidence level for each measure.",
    "Analytics Endpoints And Modules": "The endpoint grouping shows that KPI and reporting capabilities are API-backed, not only dashboard screenshots or static content.",
    "Grafana Dashboard Capabilities": "The dashboards provide operational visibility and review workflows. Formal reporting still depends on validated source data and query semantics.",
    "Node-RED Ingestion Pipeline": "The ingestion flow is the operational link between telemetry producers and the database; failures here affect every downstream analytic output.",
    "Report Generation Capabilities": "The report paths demonstrate implemented PDF generation capability while preserving the distinction between operational reports and formally audited statements.",
    "Report Workflow": "This workflow helps operators and maintainers troubleshoot report generation by separating API, data, template, rendering, and download responsibilities.",
    "Data-Quality Assumptions": "These assumptions should be treated as operational controls. When they are not met, KPI and report outputs can be technically generated but less meaningful.",
    "Audit-Use Cautions": "The cautions support responsible use of the system outputs. They define the additional governance needed for formal assurance contexts.",
    "Implemented, Configured, Partial, And Demo Data Distinctions": "These distinctions help readers understand what is ready for evaluation, what is configured for visualization, and what needs validation before production or audit use.",
    "Deployment Prerequisites": "Meeting these prerequisites reduces avoidable installation failures and gives operators the minimum environment needed for a clean first run.",
    "Docker Compose Services": "The service inventory is also the troubleshooting map: when a function fails, operators can locate the owning service and supporting dependencies.",
    "Networks, Volumes, And Ports": "These resources define the main persistence and exposure points. They require careful handling during backup, cleanup, and public deployment.",
    "Environment Variable Groups": "The grouped variables make configuration review manageable while keeping sensitive deployed values out of stakeholder documentation.",
    "Step-By-Step Deployment": "Following this sequence keeps installation reproducible and gives operators clear checkpoints before exposing the system to users.",
    "Startup, Shutdown, And Clean Reinstall Procedures": "The procedure split protects persistent data. Routine restarts should preserve volumes, while clean reinstalls require deliberate data-removal decisions.",
    "Verification Scripts And Health Checks": "The verification model separates configuration validity from live service readiness. Both are needed for a credible deployment handover.",
    "Healthcheck Details": "Healthchecks are early indicators, not full business validation. They should be combined with data freshness, dashboard, API, and workflow checks.",
    "Backup And Recovery": "The backup view identifies where durable operational state lives. Production use should include tested restoration, not only backup creation.",
    "Upgrade And Redeployment": "The redeployment sequence reduces risk by pairing source changes with backup, rebuild, smoke-test, and rollback planning.",
    "Troubleshooting Commands": "These commands let operators inspect runtime state without deleting data or changing configuration.",
    "Production Hardening Checklist": "The hardening items convert the repository's configuration hooks into production operating controls.",
    "Troubleshooting Scenarios": "The scenarios help operators start diagnosis in the correct service area rather than treating the stack as a single opaque process.",
    "Handover Summary": "The document set is complementary: architecture explains boundaries, design explains implementation, deployment explains operation, and final documentation ties the delivery together.",
    "Delivery Artifact List": "The artifact list confirms that both final DOCX deliverables and maintainable source material are part of the handover.",
    "Quick-Start Summary": "The quick-start provides orientation only. Operators should still use the deployment report for detailed configuration, hardening, and recovery decisions.",
    "Installation And Access": "Access endpoints should be verified after startup and adjusted for DNS, TLS, firewall, and target-host configuration in production.",
    "Main Workflows": "The workflows show how data and users move through the system, which is useful for demonstrations, support planning, and acceptance testing.",
    "Analytics, Dashboards, Reports, And Assistants": "The table clarifies that dashboards, API views, Rasa help, and OVOS operational queries serve different user needs and have different validation requirements.",
    "Maintenance And Troubleshooting": "These responsibilities form the ongoing operating model after initial installation.",
    "Acceptance Checklist": "The checklist confirms document completeness and traceability. Live runtime acceptance still depends on checks performed against the target deployment.",
    "Demo-Readiness Checklist": "These checks help ensure a demonstration uses current data, healthy services, and verified assistant behavior where OVOS is included.",
}


def table_intro_for(title: str) -> str:
    return TABLE_INTROS.get(
        title,
        "The table below organizes the key details for this section so readers can connect the documented capability to its operational meaning.",
    )


def table_followup_for(title: str) -> str:
    return TABLE_FOLLOWUPS.get(
        title,
        "Operationally, this table should be read as a concise evidence-backed summary. The surrounding narrative explains how the listed items affect deployment, operation, or validation.",
    )


def section_evidence(rows: Sequence[Sequence[str]]) -> Section:
    return Section(
        "Source References",
        paragraphs=[
            "The table below lists the main source material used for this document. It is not a full file inventory; it identifies the sources behind the material claims."
        ],
        table=(["Topic", "Source material"], rows),
    )


def common_limitations_section() -> Section:
    return Section(
        "Limitations And Assumptions",
        paragraphs=[
            "This section summarizes the current validation status, scope boundaries, and operational considerations for the delivered system."
        ],
        table=(["Item", "Status"], LIMITATION_ROWS),
        table_followup=[
            "Together, these points define the verified scope of the current delivery and the operational responsibilities required before production use. They preserve a clear distinction between implemented capability, deployment configuration, and assurance activities that belong to the target operating environment."
        ],
    )


def build_system_architecture(diagrams: dict) -> DocSpec:
    return DocSpec(
        filename="System Architecture Report.docx",
        title="System Architecture Report",
        purpose="Describe the implemented HumanEnerDIA / EnMS architecture and the OVOS-EnMS integration boundary.",
        audience="Project managers, technical reviewers, deployment stakeholders, and external partners.",
        evidence_note="Claims are tied to tracked source code, Docker configuration, SQL initialization files, and verified compose validation.",
        sections=[
            Section(
                "Executive Summary",
                paragraphs=[
                    "HumanEnerDIA is implemented as a Docker Compose based industrial energy management stack. It combines an Nginx gateway, a static portal, FastAPI analytics APIs, PostgreSQL/TimescaleDB storage, MQTT telemetry, Node-RED ingestion, Grafana dashboards, a simulator, authentication services, a Rasa text chatbot path, and an optional OVOS-EnMS voice/natural-language assistant layer.",
                    "The OVOS-EnMS component is a separate assistant runtime that connects to the HumanEnerDIA-compatible analytics API. Its REST bridge does not calculate energy answers by itself; it forwards user queries to the OVOS messagebus, where the EnMS skill parses, validates, executes API calls, and formats responses.",
                    ("Observed in code/config: ", "The GitHub production base stack is defined in docker-compose.yml and does not include an OVOS service. The separate OVOS-EnMS repository provides the OVOS Compose file, bridge, skill, parser, validator, API client, and response formatting source."),
                ],
            ),
            Section(
                "System Context",
                figure=("system-context.png", "Figure 1. System context and product boundaries."),
                paragraphs=[
                    "The system has three boundaries that must remain distinct in delivery documentation. HumanEnerDIA / EnMS is the energy management backend and visualization stack. OVOS-EnMS is the voice/natural-language assistant layer that integrates with the analytics API. The Rasa chatbot is a text-oriented help and knowledge path, not the same runtime as the OVOS skill.",
                ],
                table=(
                    ["Boundary", "Included components", "Evidence"],
                    [
                        ["HumanEnerDIA / EnMS", "Nginx, portal, analytics, PostgreSQL/TimescaleDB, MQTT, Node-RED, Grafana, simulator, auth-service, Rasa/chatbot services", EVIDENCE["compose"]],
                        ["OVOS-EnMS", "OVOS runtime, REST bridge, messagebus, EnMS skill, parser, validator, API client, response formatter", EVIDENCE["ovos_skill"]],
                        ["External users/integrators", "Browser users, API clients, OVOS clients, operators, WASABI reviewers", "README.md; verify.sh; docs/final-delivery/"],
                    ],
                ),
            ),
            Section(
                "Component Responsibilities",
                paragraphs=[
                    "The architecture separates gateway, domain API, storage, ingestion, visualization, simulation, authentication, text-help chatbot, and voice/natural-language assistant concerns. This separation is visible in the Compose service list and in the route/service/module layout."
                ],
                table=(["Component", "Responsibility", "Source basis"], COMPONENT_RESPONSIBILITY_ROWS),
            ),
            Section(
                "Runtime Services",
                paragraphs=[
                    "The base runtime service inventory below is taken from docker-compose.yml and validated with docker compose config --quiet for the current delivery state."
                ],
                table=(["Service", "Image or build context", "External port/path", "Responsibility", "Healthcheck"], SERVICE_ROWS),
            ),
            Section(
                "Dependency And Request Flow",
                paragraphs=[
                    "The main runtime dependencies are directional: clients enter through gateway or assistant bridge paths, domain services call databases and supporting middleware, and telemetry flows from producers into storage before it is consumed by dashboards, APIs, reports, and assistants."
                ],
                table=(["Flow", "Dependency chain"], DEPENDENCY_FLOW_ROWS),
            ),
            Section(
                "Data And Message Flow",
                figure=("telemetry-data-flow.png", "Figure 2. Telemetry ingestion and analytics data flow."),
                paragraphs=[
                    "Synthetic factory data or external device data enters through MQTT. Node-RED subscribes to factory/#, parses the topic structure, routes by payload type, validates required fields, and writes energy, production, environmental, and status data into PostgreSQL.",
                    "TimescaleDB hypertables and continuous aggregates provide raw and aggregated time-series views. The analytics service reads from those tables and aggregate views to support KPIs, baselines, forecasts, anomalies, reports, Grafana dashboards, and OVOS-facing responses.",
                ],
                table=(
                    ["Stage", "Observed implementation", "Evidence"],
                    [
                        ["Telemetry source", "simulator loads active machines from PostgreSQL and publishes MQTT messages for energy, production, environmental, status, and multi-energy boiler topics", EVIDENCE["simulator"]],
                        ["Broker", "mqtt service exposes 1883 and websocket 9001 with credentials supplied through environment variables", EVIDENCE["compose"]],
                        ["Ingestion", "Node-RED flow includes Subscribe: factory/#, Parse Topic, Route by Type, Process Energy/Production/Environmental/Status, and PostgreSQL output nodes", EVIDENCE["nodered"]],
                        ["Storage", "energy_readings, production_data, and environmental_data are converted to TimescaleDB hypertables with continuous aggregates", EVIDENCE["database_schema"]],
                        ["Consumption", "Analytics API, Grafana dashboards, portal, chatbot, and OVOS integration consume database-backed data", EVIDENCE["analytics_main"]],
                    ],
                ),
            ),
            Section(
                "External And Internal Interfaces",
                paragraphs=[
                    "External browser access normally enters through Nginx. Direct service ports are exposed for operations and development; production exposure should be restricted by firewall or reverse proxy policy."
                ],
                table=(["Interface", "Route or endpoint", "Evidence/notes"], ACCESS_ROWS + [
                    ["Analytics health", "Direct service path /api/v1/health; through Nginx analytics proxy /api/analytics/api/v1/health", "analytics/main.py; nginx/conf.d/default.conf"],
                    ["OVOS proxy via EnMS", "/api/ovos/* -> /api/v1/ovos/*", "nginx/conf.d/default.conf; analytics/api/routes/ovos_voice.py"],
                    ["OVOS direct bridge", "POST /query, POST /query/voice, GET /health", EVIDENCE["ovos_bridge"]],
                ]),
            ),
            Section(
                "OVOS-EnMS Integration",
                figure=("ovos-query-lifecycle.png", "Figure 3. OVOS request and response lifecycle."),
                paragraphs=[
                    "The OVOS bridge receives text queries through /query or /query/voice. It emits recognizer_loop:utterance to the OVOS messagebus and listens for speak and enms.skill.response events. The skill handles intent routing, context, validation, backend API calls, and deterministic response formatting.",
                    "HumanEnerDIA also exposes /api/v1/ovos/voice/query and /api/v1/ovos/voice/health as a proxy route from the analytics service to the OVOS bridge. This supports portal-side integration without making the portal responsible for OVOS messagebus details.",
                ],
            ),
            Section(
                "Deployment Variants And Boundaries",
                paragraphs=[
                    "The GitHub production tree supports a HumanEnerDIA base deployment. OVOS-EnMS should be described as a companion runtime sourced from its own repository unless and until a production Compose overlay is tracked in the HumanEnerDIA production repository."
                ],
                table=(["Variant", "Source", "Boundary statement"], DEPLOYMENT_VARIANT_ROWS),
            ),
            Section(
                "Security And Network Boundaries",
                paragraphs=[
                    "The repository supports several operational controls, but public production hardening remains an operator responsibility. The setup helper creates .env from .env.example when needed and generates first-run secrets for database, Grafana, Node-RED, Redis, MQTT, JWT, and API key values.",
                    "Authentication is implemented by auth-service using bcrypt password hashing, JWT sessions, email verification and password reset flows, admin allow-listing from environment variables, session tracking, and audit tables. Node-RED has admin authentication configured through environment-provided credentials.",
                ],
                bullets=[
                    "Do not commit .env, generated secrets, runtime logs, database dumps, model caches, or Docker volumes.",
                    "Restrict direct exposure of PostgreSQL, Redis, MQTT, Grafana, Node-RED, and service debug ports in production.",
                    "Terminate TLS at Nginx or an upstream reverse proxy before internet-facing deployment.",
                    "Rotate credentials before public use, especially any generated first-run secrets.",
                ],
            ),
            Section(
                "Operational Risks",
                paragraphs=[
                    "The following risks are not defects in the documentation package; they are deployment and governance points that should remain visible in stakeholder handover material."
                ],
                table=(["Risk area", "Observed source basis", "Operational action"], OPERATIONAL_RISK_ROWS),
            ),
            common_limitations_section(),
            section_evidence([
                ["Runtime topology", EVIDENCE["compose"]],
                ["OVOS deployment boundary", EVIDENCE["ovos_config"]],
                ["Routing", EVIDENCE["nginx"]],
                ["Database and KPIs", EVIDENCE["database_schema"]],
                ["Analytics API", EVIDENCE["analytics_main"]],
                ["OVOS bridge and skill", f"{EVIDENCE['ovos_bridge']}; {EVIDENCE['ovos_skill']}"],
                ["Compose validation", "docker compose config --quiet returned success for the HumanEnerDIA production tree and the OVOS-EnMS repository compose file"],
            ]),
        ],
    )


def build_software_design() -> DocSpec:
    return DocSpec(
        filename="Software Design Documentation.docx",
        title="Software Design Documentation",
        purpose="Document the implemented software modules, interfaces, data model, and design constraints.",
        audience="Developers, maintainers, technical reviewers, and integration engineers.",
        evidence_note="Evidence is based on route registration, service code, SQL schema, compose files, and tests rather than README-level descriptions alone.",
        sections=[
            Section(
                "Design Overview",
                paragraphs=[
                    "HumanEnerDIA uses a service-oriented design. Nginx centralizes browser and API routing; analytics owns most domain APIs; PostgreSQL/TimescaleDB owns persistent operational and time-series data; MQTT and Node-RED connect telemetry ingestion; Grafana presents dashboards; simulator produces demo telemetry; auth-service owns user/account workflows; Rasa/chatbot provides a text help assistant; OVOS-EnMS provides a separate assistant layer.",
                    "The repository favors explicit route modules and service modules rather than a single monolithic backend. The analytics service mounts routers for baselines, anomalies, KPIs, machines, forecasts, time series, visualization data, model performance, production, SEU/ISO 50001 features, reports, and OVOS-facing integration.",
                ],
            ),
            Section(
                "Module Responsibility Matrix",
                table=(
                    ["Subsystem", "Responsibilities", "Primary evidence"],
                    [
                        ["analytics/api/routes", "FastAPI request handlers and route-specific request/response behavior", EVIDENCE["analytics_routes"]],
                        ["analytics/services", "Business logic for KPIs, baselines, forecasts, anomaly handling, performance, event publishing, reports, and Redis coordination", EVIDENCE["analytics_services"]],
                        ["analytics/models", "ML/statistical model implementations and model persistence helpers", EVIDENCE["analytics_models"]],
                        ["database/init", "First-start schema, hypertables, continuous aggregates, SQL functions, seed data, ISO 50001 and model-performance tables", EVIDENCE["database_schema"]],
                        ["simulator", "FastAPI control endpoints, machine simulation classes, MQTT publisher, auto anomaly injection support", EVIDENCE["simulator"]],
                        ["nodered", "MQTT topic parsing, data validation, and PostgreSQL write pipeline", EVIDENCE["nodered"]],
                        ["auth-service", "Registration, login, JWT verification, admin APIs, email verification/reset, pilot/contact forms", EVIDENCE["auth"]],
                        ["chatbot/rasa", "Text help chatbot, QA retrieval actions, Rasa runtime, Express proxy backend", EVIDENCE["chatbot"]],
                    ],
                ),
            ),
            Section(
                "Analytics Service Design",
                paragraphs=[
                    "The analytics service is a FastAPI application with lifespan-managed database connection, optional Redis event subscriber, scheduler startup, route registration, CORS middleware, request logging, timeout handling, and generic exception handling.",
                    "Router registration in analytics/main.py shows the implemented surface: baseline, anomaly, KPI, machines, forecast, time series, sankey, heatmap, comparison, model performance, stats, production, SEU/factory/performance/ISO 50001/multi-energy, OVOS, OVOS voice proxy, and reports.",
                ],
                table=(["API area", "Representative endpoints", "Evidence"], API_ROUTE_ROWS),
            ),
            Section(
                "API Design Details",
                paragraphs=[
                    "The API design is organized by domain routers rather than by a single generic endpoint. The production tree exposes both operational APIs and UI-serving routes, and some compatibility routes remain mounted for OVOS-oriented integrations."
                ],
                table=(["Design aspect", "Observed implementation", "Source basis"], API_DESIGN_DETAIL_ROWS),
            ),
            Section(
                "Database And Schema Design",
                paragraphs=[
                    "The database initialization files create core dimensions, time-series facts, current-state tables, baseline/anomaly/tariff/carbon/audit tables, auth tables, ISO 50001 tables, model-performance tables, forecast output tables, and action-plan workflow tables.",
                    "TimescaleDB is used for high-frequency time-series storage. The initialization scripts create hypertables for energy_readings, production_data, environmental_data, and energy_forecasts, plus continuous aggregates at 1 minute, 15 minutes, 1 hour, and 1 day where implemented.",
                ],
                table=(
                    ["Database object group", "Implemented objects", "Evidence"],
                    [
                        ["Core entities", "factories, machines, energy_readings, production_data, environmental_data, machine_status", "database/init/02-schema.sql"],
                        ["Analytics metadata", "energy_baselines, anomalies, energy_tariffs, carbon_factors, model performance/training/alert tables", "database/init/02-schema.sql; 11-13 model scripts"],
                        ["ISO 50001", "energy_sources, seus, seu_energy_performance, enpi_baselines, enpi_performance, energy_targets, action_plans", "database/init/07-iso50001-schema.sql; 15-16 scripts"],
                        ["Aggregates", "energy, production, and environmental aggregate materialized views", "database/init/03-timescaledb-setup.sql"],
                        ["KPI functions", "calculate_sec, calculate_peak_demand, calculate_load_factor, calculate_energy_cost, calculate_carbon_intensity, calculate_all_kpis", "database/init/04-functions.sql"],
                    ],
                ),
            ),
            Section(
                "Database Object Catalog",
                paragraphs=[
                    "The following catalog is intended as a stakeholder-level data-design view. It avoids column-by-column schema reproduction while identifying the functional database groups that support the application."
                ],
                table=(["Object group", "Representative objects", "Design role"], DATABASE_DESIGN_ROWS),
            ),
            Section(
                "Service-Layer Design",
                paragraphs=[
                    "The analytics service layer centralizes most non-trivial domain logic. KPIService wraps SQL KPI functions; baseline and forecast services coordinate model training/prediction and storage; anomaly services detect and record anomalies; report services assemble data and output files; event publisher/subscriber modules integrate Redis Pub/Sub where enabled.",
                    "This design keeps route handlers closer to request/response orchestration while delegating calculation, modeling, report assembly, and event behavior to domain services. Some older routes still contain direct SQL or route-level calculations, so the design is pragmatic rather than fully uniform."
                ],
                table=(
                    ["Service area", "Implementation responsibility", "Source basis"],
                    [
                        ["KPI service", "Calls SQL functions for SEC, peak demand, load factor, cost, carbon, and combined KPI responses.", "analytics/services/kpi_service.py; database/init/04-functions.sql"],
                        ["Baseline services", "Train/predict baseline models, store model metadata, explain drivers, and support SEU baseline training.", "analytics/services/baseline_service.py; seu_baseline_service.py; analytics/models/baseline.py"],
                        ["Forecast service", "Train ARIMA/Prophet models, create predictions, and support short-term/peak/demand routes.", "analytics/services/forecast_service.py; analytics/models/*forecast*.py"],
                        ["Anomaly service", "Create/detect/search/resolve anomalies and support anomaly-oriented routes.", "analytics/services/anomaly_service.py; analytics/api/routes/anomaly.py"],
                        ["Report services", "Generate legacy monthly EnPI PDFs and V2 report outputs through report components and generators.", "analytics/reports/; analytics/reports_v2/"],
                        ["Event services", "Publish and subscribe to Redis channels for anomaly, metric, training, and system alert events when enabled.", "analytics/services/event_publisher.py; event_subscriber.py; redis_manager.py"],
                    ],
                ),
            ),
            Section(
                "Simulator And Ingestion Design",
                paragraphs=[
                    "The simulator is a FastAPI service with lifecycle initialization. It connects to PostgreSQL, connects to MQTT, loads active machines from the database, creates simulator instances by machine type, and can auto-start based on configuration.",
                    "Machine implementations generate energy, production, environmental, and status payloads. The boiler path supports multi-energy publication for electricity, natural gas, and steam style payloads. Node-RED processes subscribed MQTT traffic and writes normalized records into the database.",
                ],
                table=(
                    ["Area", "Design details", "Evidence"],
                    [
                        ["Control API", "start, stop, runtime config, status, list machines, machine detail, inject/clear anomaly, info", "simulator/api/routes.py"],
                        ["Machine loading", "Loads active machines from database with type, rated_power_kw, interval, and MQTT topic", "simulator/simulator_manager.py"],
                        ["MQTT publishing", "Publishes energy, multi-energy, production, environmental, and retained status messages", "simulator/mqtt_publisher.py"],
                        ["Node-RED flow", "Subscribe: factory/#, Parse Topic, Route by Type, Process Energy/Production/Environmental/Status", "nodered/data/flows.json"],
                    ],
                ),
            ),
            Section(
                "Authentication, Portal, And Chatbot Design",
                paragraphs=[
                    "auth-service is a Flask application backed by demo_users, demo_sessions, demo_audit_log, and pilot_factory_applications tables. It implements registration, login, JWT verification, email verification, password reset, admin user management, CSV export, pilot factory application workflows, and contact form handling.",
                    "The portal is static HTML/CSS/JS served by Nginx. It includes general pages, authentication pages, admin pages, report pages, and an OVOS voice widget script. The chatbot backend is an Express service that serves the built frontend and proxies to Rasa and OVOS endpoints.",
                    "The Rasa custom action loads qa_data.json and retrieves knowledge/help answers using exact match, special cases, keyword routing, abbreviation expansion, misspelling correction, and fuzzy-style matching logic. This text help path is separate from live OVOS operational queries.",
                ],
            ),
            Section(
                "Authentication And Authorization Design",
                paragraphs=[
                    "Authentication is implemented in a separate Flask service rather than inside the analytics FastAPI application. The service stores users, sessions, and audit records in PostgreSQL auth tables and exposes login/registration/admin/contact workflows through Nginx-routed endpoints.",
                    "Authorization is strongest on the auth-service admin endpoints, where a decorator verifies bearer JWTs, checks the configured admin email allowlist, and confirms the active/verified admin role in the database. Analytics middleware includes JWT support; the implemented scope is best described as service-level authentication and admin authorization rather than a complete cross-service enterprise IAM layer."
                ],
                table=(
                    ["Auth feature", "Observed implementation", "Source basis"],
                    [
                        ["Password storage", "bcrypt with 12 rounds.", "auth-service/auth_service.py"],
                        ["Session token", "JWT signed with HS256 and stored in demo_sessions with expiry metadata.", "auth-service/auth_service.py; database/init/05-auth-schema.sql"],
                        ["Email verification", "Verification tokens and verified_at fields; email can be disabled, which auto-verifies users on registration.", "auth-service/auth_service.py"],
                        ["Password reset", "Reset token and timestamp fields with one-hour expiry logic.", "auth-service/auth_service.py; database/init/05-auth-schema.sql"],
                        ["Admin controls", "Admin allowlist from ADMIN_EMAILS plus database role/is_active/email_verified checks.", "auth-service/auth_service.py"],
                        ["Audit trail", "REGISTER, LOGIN, EMAIL_VERIFY and related actions are inserted into demo_audit_log.", "auth-service/auth_service.py; database/init/05-auth-schema.sql"],
                    ],
                ),
            ),
            Section(
                "Configuration, Validation, Logging, And Error Handling",
                bullets=[
                    "Configuration is primarily environment-driven through .env.example, docker-compose.yml, analytics/config.py, simulator/config.py, Node-RED settings, and OVOS settings/config files.",
                    "The setup helper preserves existing non-placeholder .env values, generates missing first-run secrets, validates Compose, builds images, and starts services.",
                    "FastAPI services use health endpoints, request logging, validation exception handlers, and generic exception handlers.",
                    "OVOS skill validation uses Pydantic schemas, confidence thresholding, machine whitelists, fuzzy matching, metric validation, time-range parsing, and entity normalization.",
                    "auth-service uses bcrypt password hashing, JWT sessions, email verification gates, admin decorators, and parameterized SQL queries.",
                ],
            ),
            Section(
                "Configuration Reference",
                paragraphs=[
                    "Configuration is intentionally environment-driven. The delivery documentation describes configuration groups and operational responsibility without disclosing actual .env values."
                ],
                table=(["Configuration group", "Representative variables", "Source basis"], CONFIG_GROUP_ROWS),
                table_followup=[
                    "These configuration groups describe the main runtime controls across database access, service ports, security, telemetry, analytics behavior, and optional OVOS proxying. Deployment owners should manage the actual values in the target environment."
                ],
            ),
            Section(
                "Error Handling And Logging",
                paragraphs=[
                    "The codebase includes explicit error handling in key services, but behavior is not completely uniform across every route. This should be represented as implemented service-level error handling rather than a single formal enterprise error contract."
                ],
                table=(["Area", "Observed behavior"], ERROR_HANDLING_ROWS),
            ),
            Section(
                "Known Design Gaps And Placeholders",
                table=(["Gap or caution", "Source-backed status"], [
                    ["Report V2 semantics", "V2 routes and generator exist, but some service values are proportional or placeholder-derived, such as efficiency sparkline and estimated baseline cost."],
                    ["Simulator machine list inconsistency", "Code supports boiler; simulator info endpoint text still lists five machine types."],
                    ["Direct public exposure", "Several internal service ports are externally mapped for development/ops; production hardening requires operator firewall/TLS review."],
                    ["README claims", "Root README contains high-level feature claims; final documents use code/config evidence where details differ."],
                ]),
            ),
            section_evidence([
                ["Analytics app and routers", f"{EVIDENCE['analytics_main']}; {EVIDENCE['analytics_routes']}"],
                ["SQL schema/functions", EVIDENCE["database_schema"]],
                ["Simulator", EVIDENCE["simulator"]],
                ["Node-RED", EVIDENCE["nodered"]],
                ["Auth", EVIDENCE["auth"]],
                ["Chatbot/Rasa", EVIDENCE["chatbot"]],
                ["Validation performed", "docker compose config --quiet for HumanEnerDIA production; docker compose -f <OVOS-EnMS repository>/docker-compose.yml config --quiet"],
            ]),
        ],
    )


def build_skill_doc(diagrams: dict) -> DocSpec:
    return DocSpec(
        filename="Skill Documentation.docx",
        title="Skill Documentation",
        purpose="Document the OVOS-EnMS skill, REST bridge, parser, validation, API client, and response behavior.",
        audience="OVOS integrators, WASABI technical reviewers, backend maintainers, and external partners.",
        evidence_note="OVOS-EnMS evidence comes from the separate OVOS-EnMS source repository, with HumanEnerDIA API integration evidence from the HumanEnerDIA production tree.",
        sections=[
            Section(
                "Purpose And Boundaries",
                paragraphs=[
                    "The HumanEnerDIA OVOS skill is the natural-language assistant layer for industrial energy-management questions. It is not the HumanEnerDIA backend and does not own telemetry storage or KPI calculation. It connects to a reachable HumanEnerDIA-compatible analytics API.",
                    "The production integration boundary is the HumanEnerDIA-compatible REST API. The repository includes adapter abstractions, but v1.0.0 documentation states that arbitrary third-party EnMS APIs require an adapter or proxy that exposes the expected API contract.",
                ],
            ),
            Section(
                "Deployment And Configuration",
                paragraphs=[
                    "The separate OVOS-EnMS repository provides a Docker Compose service that exposes the REST bridge on port 5000 and the OVOS messagebus on port 8181. The HumanEnerDIA production base docker-compose.yml does not define an OVOS service, so production documentation must treat OVOS-EnMS as a companion repository/runtime unless a production overlay is explicitly added and tracked.",
                    "Key configuration includes ENMS_API_URL, OVOS_BRIDGE_PORT, STRUCTURED_RESPONSE_GRACE_SECONDS, OVOS_TTS_ENABLED, LOG_LEVEL, OVOS_CONFIG_PATH, and XDG_CONFIG_HOME. Skill-level settings include enms_api_base_url, llm_model_path, confidence_threshold, and progress feedback options.",
                ],
                table=(
                    ["Configuration item", "Observed default or behavior", "Evidence"],
                    [
                        ["ENMS_API_URL", "Docker default points at a HumanEnerDIA-compatible /api/v1 backend", "OVOS-EnMS repository: docker-compose.yml"],
                        ["enms_api_base_url", "Skill setting for backend API URL", "settings.docker.json; settingsmeta.yaml"],
                        ["confidence_threshold", "Default 0.85 in settings and validator configuration", "settings.docker.json; lib/validator.py"],
                        ["INSTALL_LLM_FALLBACK", "Build argument for installing optional LLM dependencies in the Dockerfile", "OVOS-EnMS repository: Dockerfile"],
                    ],
                ),
            ),
            Section(
                "Configuration Reference",
                paragraphs=[
                    "The configuration items below are taken from the OVOS-EnMS Dockerfile, Compose file, settings files, bridge, and validator. They are operational settings, not secrets; actual runtime values should still be reviewed in the deployed environment."
                ],
                table=(["Setting", "Location", "Purpose"], SKILL_CONFIG_REFERENCE_ROWS),
            ),
            Section(
                "Query Lifecycle",
                figure=("ovos-query-lifecycle.png", "Figure 1. REST bridge, messagebus, skill, API, and response lifecycle."),
                paragraphs=[
                    "The REST bridge exposes GET /health and POST /query. POST /query/voice is an alias used by the analytics proxy when audio-capable flows request the same bridge behavior.",
                    "For each query, the bridge creates or uses a session id, emits recognizer_loop:utterance to the OVOS messagebus, and waits for a speak message plus, when available, an enms.skill.response structured payload. The response returns success status, spoken response text, intent, confidence, data, insights, timestamp, and session id.",
                    "The EnMS skill receives the utterance through OVOS intent handlers or fallback handling. It parses the utterance, validates intent/entity output, calls the configured backend API, formats a deterministic response, speaks it, and emits structured response data for the bridge or portal widget.",
                ],
            ),
            Section(
                "Supported Intent And Query Families",
                paragraphs=[
                    "The active IntentType enum and skill handlers show the supported query families below. This table is not a guarantee that every phrasing is understood; it identifies implemented categories in the skill code."
                ],
                table=(["Intent family", "Purpose"], INTENT_ROWS),
            ),
            Section(
                "Intent Parsing And Routing",
                paragraphs=[
                    "The parser is hybrid. Tier 1 is regex-based heuristic routing for common operational queries. Tier 2 uses Adapt pattern matching and registered vocabulary. Tier 3 is an optional local Qwen GGUF LLM parser used as fallback when dependencies and model files are available.",
                    "The active parser code includes patterns for production, anomaly detection, forecasts, KPIs, performance, baselines, driver analysis, SEUs, rankings, factory overview, status, power, and related query types. Adapt vocabulary registers machine names, spoken number variants, energy/power/status/cost/KPI/factory/comparison/time/forecast/anomaly/help terms and more.",
                ],
                table=(
                    ["Tier", "Implementation", "Important note"],
                    [
                        ["Heuristic", "Regex patterns in lib/intent_parser.py", "Fast path for common operational wording."],
                        ["Adapt", "IntentDeterminationEngine in lib/adapt_parser.py", "Pattern/vocabulary matching with registered machine and domain terms."],
                        ["LLM", "Qwen3Parser in lib/llm_parser.py", "Optional fallback requiring llama-cpp-python and a GGUF model file."],
                    ],
                ),
            ),
            Section(
                "Validation And Fuzzy Matching",
                paragraphs=[
                    "Validation is deliberately conservative. The validator builds a Pydantic Intent model, checks confidence, rejects unknown intent types, validates machine names against a whitelist, supports fuzzy matching and number-word normalization, detects ambiguity, validates multi-machine comparisons, and performs soft metric validation.",
                    "Machine discovery can refresh the whitelist from the backend API during runtime; fallback machine names are configured for cases where API discovery fails. This helps prevent hallucinated machine names from becoming backend calls.",
                ],
            ),
            Section(
                "Backend API Client And Adapter Behavior",
                paragraphs=[
                    "The ENMSClient wraps async HTTP calls to the configured backend. It uses connection pooling, request timeout management, and tenacity retry behavior that retries connection/timeouts and server-side 5xx responses while avoiding retries on ordinary 4xx client errors.",
                    "Client methods cover health, stats, machines, time series, top consumers, anomalies, KPIs, performance opportunities, action plans, forecasts, baseline models/explanations, SEU/energy-source data, reports, and ISO 50001 EnPI/action-plan endpoints. The production path remains HumanEnerDIA-compatible API usage.",
                ],
                table=(
                    ["Client area", "Representative methods", "Evidence"],
                    [
                        ["System and machines", "health_check, system_stats, factory_summary, list_machines, get_machine_status", EVIDENCE["ovos_client"]],
                        ["Telemetry", "get_energy_timeseries, get_power_timeseries, get_latest_reading, get_multi_machine_energy", EVIDENCE["ovos_client"]],
                        ["Analytics", "detect_anomalies, get_all_kpis, analyze_performance, forecast_demand, predict_baseline", EVIDENCE["ovos_client"]],
                        ["Reports and ISO", "get_enpi_report, list_action_plans, get_report_types, preview_report, generate_report", EVIDENCE["ovos_client"]],
                    ],
                ),
            ),
            Section(
                "Backend Method Mapping",
                paragraphs=[
                    "This mapping connects natural-language intent families to backend client methods and the HumanEnerDIA-compatible API areas they use. It defines implementation traceability; accepted phrasing still depends on parser coverage, validation, backend data, and deployment health."
                ],
                table=(["Intent family", "Primary backend method(s)", "API area", "Returned information"], SKILL_BACKEND_MAPPING_ROWS),
            ),
            Section(
                "Response Formatting",
                paragraphs=[
                    "The response formatter uses Jinja2 templates and custom number/unit/time filters. The formatter documentation and code explicitly state that final responses should come from API data and templates rather than free-form LLM generation.",
                    "Additional enrichment exists for anomaly responses, including severity grouping, resolved/unresolved counts, metric/anomaly label humanization, and concise spoken examples.",
                ],
            ),
            Section(
                "Example Supported Queries",
                paragraphs=[
                    "The examples below are representative of implemented intent categories and handler/parser coverage. Exact results depend on available machines, backend data, current telemetry, and deployment health."
                ],
                table=(["Category", "Example query"], SKILL_EXAMPLE_ROWS),
            ),
            Section(
                "Optional LLM Fallback",
                paragraphs=[
                    "Fast heuristic and Adapt routing are the normal path. The OVOS-EnMS Dockerfile installs LLM fallback dependencies only when INSTALL_LLM_FALLBACK=true, and the skill settings point to a configurable GGUF model path. Model-file availability must be verified in the OVOS-EnMS runtime; the HumanEnerDIA production Compose file does not bundle or start OVOS.",
                    "The LLM parser uses llama-cpp-python when installed, loads a configured GGUF model, performs deterministic JSON intent classification, and returns None on missing dependencies, missing model, parse failures, or timeout. It should be documented as optional fallback, not as required normal operation.",
                ],
            ),
            Section(
                "Failure Behavior",
                paragraphs=[
                    "The assistant layer should be represented as conservative. Parser/validator failures, missing machines, backend errors, and missing optional LLM dependencies should result in clarification or failure responses rather than fabricated energy data."
                ],
                table=(["Failure case", "Observed behavior"], SKILL_FAILURE_ROWS),
            ),
            common_limitations_section(),
            section_evidence([
                ["REST bridge", EVIDENCE["ovos_bridge"]],
                ["Skill lifecycle and handlers", EVIDENCE["ovos_skill"]],
                ["Intent parser tiers", EVIDENCE["ovos_parser"]],
                ["Validation", EVIDENCE["ovos_validator"]],
                ["API client", EVIDENCE["ovos_client"]],
                ["Response formatter", EVIDENCE["ovos_formatter"]],
                ["Configuration and deployment", EVIDENCE["ovos_config"]],
            ]),
        ],
    )


def build_kpi_doc(diagrams: dict) -> DocSpec:
    return DocSpec(
        filename="Energy Management System Reports and KPI Reports.docx",
        title="Energy Management System Reports and KPI Reports",
        purpose="Document implemented energy data, KPI, dashboard, and report capabilities with formulas and evidence.",
        audience="Energy managers, project reviewers, operators, analytics maintainers, and external partners.",
        evidence_note="KPI formulas are included only where defined in SQL functions, code, dashboard queries, or tracked documentation.",
        sections=[
            Section(
                "Energy Data Model",
                figure=("telemetry-data-flow.png", "Figure 1. Data path from telemetry to KPI/report consumers."),
                paragraphs=[
                    "HumanEnerDIA stores factory/site records, machines/SEUs, high-frequency energy readings, production data, environmental context, machine status, baseline metadata, anomaly records, tariffs, carbon factors, audit records, ISO 50001 entities, model tracking, forecast output, and action plans.",
                    "The first-start seed data defines two sample factories and eight sample machines across the demo and European facilities. The machine examples include Compressor-1, HVAC-Main, Conveyor-A, Hydraulic-Pump-1, Injection-Molding-1, Boiler-1, Compressor-EU-1, and HVAC-EU-North.",
                ],
                table=(
                    ["Concept", "Implemented representation", "Evidence"],
                    [
                        ["Factories", "factories table and seed data for Demo Manufacturing Plant and European Production Facility", "database/init/02-schema.sql; 06-seed-data.sql"],
                        ["Machines/SEUs", "machines table plus ISO-oriented seus and SEU performance tables", "database/init/02-schema.sql; 07-iso50001-schema.sql"],
                        ["Energy readings", "energy_readings hypertable with energy_type, power, energy, electrical quality fields, metadata", "database/init/02-schema.sql; 03-timescaledb-setup.sql"],
                        ["Production data", "production_data hypertable for production count, quality, throughput, mode, downtime", "database/init/02-schema.sql"],
                        ["Environmental data", "environmental_data hypertable for temperature, humidity, pressure, flow, HVAC, vibration context", "database/init/02-schema.sql"],
                        ["Energy sources", "energy_sources and energy_source_features support multi-energy/source-aware modeling", "database/init/07-iso50001-schema.sql; 10a-energy-source-features.sql"],
                    ],
                ),
            ),
            Section(
                "KPI Formula Evidence",
                paragraphs=[
                    "The following KPI formulas are implemented as database functions and wrapped by analytics/services/kpi_service.py. Some additional API routes compute aggregate factory cost/carbon estimates with constants; those should be described as route-level estimates rather than tariff/factor driven SQL functions."
                ],
                table=(["KPI", "Formula or calculation", "Implementation", "Evidence"], KPI_ROWS),
            ),
            Section(
                "KPI Catalog And Classification",
                paragraphs=[
                    "This catalog distinguishes implemented KPI formulas from configured dashboard/reporting views and route-level estimates. It is intentionally conservative so stakeholder readers can see which measures are calculation-backed and which require review before formal use."
                ],
                table=(["KPI/reporting item", "Classification", "Implementation basis", "Caution"], KPI_CATALOG_ROWS),
            ),
            Section(
                "Analytics Endpoints And Modules",
                table=(
                    ["Capability area", "Implemented routes/modules", "Notes"],
                    [
                        ["KPI", "/api/v1/kpi/sec, /factory, /factories, /peak-demand, /load-factor, /energy-cost, /carbon, /all", "Machine and factory KPI endpoints exist; formulas vary by endpoint."],
                        ["Baselines", "/baseline/train, /deviation, /predict, /models, /drivers, /train-seu", "ML baseline model metadata and saved model files are present."],
                        ["Forecasts", "/forecast/train/arima, /train/prophet, /predict, /demand, /optimal-schedule, /models, /peak, /short-term", "Uses forecasting model modules and forecast prediction tables."],
                        ["Anomalies", "/anomaly/create, /detect, /search, /recent, /active, /resolve", "Anomaly detection and search APIs with anomaly table evidence."],
                        ["Performance and ISO 50001", "/performance/analyze, /opportunities, /action-plan, /health; /iso50001/*", "Performance engine and ISO 50001 action-plan/reporting workflows are implemented."],
                        ["Production", "/production/{machine_id}", "Production metrics and related energy/cost/carbon estimates are exposed."],
                        ["Reports", "/reports/generate, /preview, /v2/generate, /v2/download/{id}, /v2/status", "Legacy monthly EnPI PDF and newer V2 report system exist."],
                    ],
                ),
            ),
            Section(
                "Grafana Dashboard Capabilities",
                paragraphs=[
                    "Grafana provisioning and dashboard JSON files are present. The dashboard inventory below is based on the tracked JSON dashboard titles and panel names. Dashboard presence is evidence of configured reporting views, while formal audit use requires validation of panel SQL, source tables, time filters, tariff/factor assumptions, and data freshness."
                ],
                table=(["Dashboard", "Panel themes"], DASHBOARD_ROWS),
            ),
            Section(
                "Dashboard Interpretation Notes",
                paragraphs=[
                    "Dashboards should be treated as operational and review views over the tracked SQL/panel configuration. They are valuable for demonstration, monitoring, and stakeholder walkthroughs, but a formal audit should trace each panel query to source tables, time filters, tariff/factor assumptions, and data freshness.",
                    "There are duplicate upper/lowercase dashboard JSON variants for several SOTA dashboards in the production tree. The documentation records the configured dashboard capability without treating duplicated JSON files as separate business capabilities."
                ],
            ),
            Section(
                "Node-RED Ingestion Pipeline",
                paragraphs=[
                    "The tracked Node-RED flow subscribes to MQTT topic factory/# and includes function nodes for topic parsing, route selection, payload validation, database preparation, success counting, error catching, and a 30-second statistics dashboard update. Credential files are intentionally not inspected or reproduced in this package.",
                ],
                table=(
                    ["Flow area", "Observed nodes", "Evidence"],
                    [
                        ["Input", "MQTT in node Subscribe: factory/#", "nodered/data/flows.json"],
                        ["Routing", "Parse Topic, Route by Type", "nodered/data/flows.json"],
                        ["Processing", "Process Energy, Process Production, Process Environmental, Process Status", "nodered/data/flows.json"],
                        ["Storage", "PostgreSQL nodes via node-red-contrib-postgresql", "nodered/package.json; nodered/data/flows.json"],
                        ["Monitoring/errors", "Count Success, Catch All Errors, Log Error, Stats Dashboard", "nodered/data/flows.json"],
                    ],
                ),
            ),
            Section(
                "Report Generation Capabilities",
                paragraphs=[
                    "The legacy report path exposes a monthly_enpi report type, generates report data, generates machine and daily trend charts, and returns a ReportLab PDF. The V2 report path creates a report id, writes a PDF under /tmp, and exposes a download endpoint. V2 components include cover page, executive dashboard, energy overview, machine analysis, cost analysis, and carbon analysis templates/components.",
                    "Important caution: the V2 generator is implemented, but some values are derived or placeholder-like in code. Examples include proportional cost/carbon trend assumptions and constant efficiency sparkline values. The final report should therefore be presented as implemented reporting capability, not as independently audited KPI methodology.",
                ],
                table=(
                    ["Report path", "Implemented behavior", "Evidence"],
                    [
                        ["Legacy monthly EnPI", "GET /types, POST /generate, GET /preview for monthly_enpi", "analytics/api/routes/reports.py; analytics/reports/monthly_enpi_report.py"],
                        ["V2 PDF report", "POST /v2/generate, GET /v2/download/{report_id}, GET /v2/status", "analytics/api/routes/reports.py; analytics/reports_v2/services/report_service.py"],
                        ["V2 templates", "Base, header/footer, KPI cards, chart containers, cover, executive dashboard, energy overview, machine ranking/profile, cost, carbon sections", "analytics/reports_v2/templates/"],
                    ],
                ),
            ),
            Section(
                "Report Workflow",
                paragraphs=[
                    "The report workflow separates API request handling, data gathering, chart/template generation, PDF creation, and download/preview behavior. Formal audit use requires validation of report semantics, formulas, source data, tariff factors, carbon factors, and period boundaries."
                ],
                table=(["Step", "Observed behavior", "Source basis"], REPORT_WORKFLOW_ROWS),
            ),
            Section(
                "Data-Quality Assumptions",
                paragraphs=[
                    "KPI and report outputs assume data freshness, correctly associated machine/factory records, synchronized timestamps, and appropriate tariff/carbon-factor records. These assumptions are not always enforceable by the codebase alone and should be part of operational handover."
                ],
                table=(["Assumption", "Reason"], DATA_QUALITY_ROWS),
            ),
            Section(
                "Audit-Use Cautions",
                paragraphs=[
                    "The system implements carbon, cost, EnPI, SEC, forecasting, anomaly, and report capabilities where code/config evidence exists. That does not automatically make every generated output suitable for regulatory, financial, or audited sustainability reporting without governance review."
                ],
                table=(["Caution", "Stakeholder guidance"], AUDIT_CAUTION_ROWS),
            ),
            Section(
                "Implemented, Configured, Partial, And Demo Data Distinctions",
                table=(
                    ["Capability", "Classification", "Reason"],
                    [
                        ["TimescaleDB energy/production/environmental storage", "Supported by implementation", "Tables, hypertables, and aggregate views are created by SQL init scripts."],
                        ["SEC, peak demand, load factor, cost, carbon KPI functions", "Supported by implementation", "SQL functions and service wrappers exist."],
                        ["Grafana dashboards", "Configured", "Dashboard JSON and provisioning are tracked."],
                        ["Node-RED ingestion", "Configured and implemented", "Flow nodes and settings are tracked; live execution requires validation on the target deployment."],
                        ["Sample factories and machines", "Demo/sample data", "Seed SQL inserts named sample facilities and machines."],
                        ["V2 report polish/semantic completeness", "Partially implemented", "Routes/templates exist, but some data calculations are placeholders or estimates."],
                        ["Standalone query API service", "Out of scope", "No separate natural-language query API service is defined in the GitHub production docker-compose.yml."],
                    ],
                ),
            ),
            common_limitations_section(),
            section_evidence([
                ["KPI functions", "database/init/04-functions.sql"],
                ["KPI routes and service", "analytics/api/routes/kpi.py; analytics/services/kpi_service.py"],
                ["Report routes/services", EVIDENCE["reports"]],
                ["Database schema", EVIDENCE["database_schema"]],
                ["Node-RED", EVIDENCE["nodered"]],
                ["Grafana", EVIDENCE["grafana"]],
                ["Simulator seed data", EVIDENCE["seed_data"]],
            ]),
        ],
    )


def build_docker_doc(diagrams: dict) -> DocSpec:
    return DocSpec(
        filename="Docker Deployment Report.docx",
        title="Docker Deployment Report",
        purpose="Document Docker Compose deployment, configuration, startup, verification, health checks, and operational troubleshooting.",
        audience="Operators, deployment engineers, technical reviewers, and external partner infrastructure teams.",
        evidence_note="Deployment claims are based on compose files, Dockerfiles, setup and verification scripts, and compose validation.",
        sections=[
            Section(
                "Deployment Overview",
                figure=("deployment-startup-flow.png", "Figure 1. Deployment preparation, setup, validation, build, start, and verification flow."),
                paragraphs=[
                    "The deployment target described by the repository is a Linux host running Docker Engine and Docker Compose v2. The GitHub production base stack is defined by docker-compose.yml. The separate OVOS-EnMS repository provides its own Compose file and should be deployed/validated separately when the assistant runtime is in scope.",
                    "The setup helper is the intended guided path. It creates .env from .env.example when needed, generates first-run secrets for placeholders, validates Docker Compose, builds images, and starts the stack. It also adjusts OVOS_BRIDGE_HOST when an optional OVOS compose file is present.",
                ],
            ),
            Section(
                "Deployment Prerequisites",
                paragraphs=[
                    "The repository does not install host Docker itself. Operators should prepare the host and confirm prerequisite tooling before running setup or Compose commands."
                ],
                table=(
                    ["Prerequisite", "Source basis / action"],
                    [
                        ["Linux host or compatible container host", "Deployment scripts assume a shell environment with Docker available."],
                        ["Docker Engine and Docker Compose v2", "setup.sh requires docker and verifies docker compose version."],
                        ["Outbound image/build access", "Compose builds local images and pulls base images such as nginx, TimescaleDB, Redis, and Grafana."],
                        ["curl and grep for verification", "verify.sh requires curl and grep."],
                        ["Optional openssl or /dev/urandom", "setup.sh uses openssl rand when available for generated secrets, otherwise /dev/urandom."],
                        ["Optional python3-bcrypt or Docker", "setup.sh needs one path to generate the Node-RED bcrypt password hash."],
                        ["Ports available", "Default external ports include 8080, 8443, 5433, 1883, 9001, 3001, 1881, 8001, 8003, 5500, 5005, 5006, 5055, 6380, 5000, and 8181."],
                    ],
                ),
            ),
            Section(
                "Compose Service Topology",
                figure=("docker-service-topology.png", "Figure 2. Docker Compose service topology and optional OVOS attachment."),
                paragraphs=[
                    "The base deployment uses one Docker bridge network for HumanEnerDIA services. Nginx is the browser/API gateway; analytics, auth-service, chatbot/Rasa, simulator, Node-RED, Grafana, PostgreSQL/TimescaleDB, MQTT, and Redis communicate on the internal network. OVOS-EnMS remains a separate assistant runtime in the tracked production source and connects through the HumanEnerDIA-compatible API."
                ],
            ),
            Section(
                "Docker Compose Services",
                table=(["Service", "Image or build context", "External port/path", "Responsibility", "Healthcheck"], SERVICE_ROWS),
            ),
            Section(
                "Networks, Volumes, And Ports",
                paragraphs=[
                    "All HumanEnerDIA production Compose services join the Docker bridge network named by ENMS_NETWORK_NAME, defaulting to enms-network. OVOS-EnMS is documented as a separate assistant runtime rather than a service in the GitHub production base docker-compose.yml.",
                    "Persistent named volumes in the production Compose file include PostgreSQL data, MQTT data/logs, Redis data, Node-RED data, and Grafana data.",
                ],
                table=(
                    ["Resource", "Configured name/default", "Purpose"],
                    [
                        ["Network", "${ENMS_NETWORK_NAME:-enms-network}", "Service-to-service communication"],
                        ["postgres-data", "${VOLUME_PREFIX:-enms}-postgres-data", "PostgreSQL/TimescaleDB persistent data"],
                        ["grafana-data", "${VOLUME_PREFIX:-enms}-grafana-data", "Grafana runtime data"],
                        ["redis-data", "${VOLUME_PREFIX:-enms}-redis-data", "Redis append-only persistence"],
                        ["mqtt-data/logs", "${VOLUME_PREFIX:-enms}-mqtt-data and -mqtt-logs", "Mosquitto runtime data and logs"],
                    ],
                ),
            ),
            Section(
                "Environment Variables And Configuration",
                paragraphs=[
                    ".env.example is the safe public configuration template. The real .env file is intentionally not included and must not be committed or copied into documentation. The setup helper generates first-run values for placeholders and preserves existing non-placeholder values.",
                    "Important configuration groups include database credentials, Redis password, MQTT credentials, Grafana admin credentials, Node-RED credential secret and password hash, JWT secret, API key, server IP/frontend URL, Grafana root URL, simulator controls, OVOS bridge host/port/timeout, and SMTP/admin settings.",
                ],
                bullets=[
                    "Use .env.example in documentation, not .env.",
                    "Rotate generated first-run credentials before production exposure.",
                    "Set DNS, TLS, firewall rules, and public URLs explicitly for production.",
                ],
            ),
            Section(
                "Environment Variable Groups",
                paragraphs=[
                    "The table below summarizes configuration groups without reproducing private runtime values. Use .env.example and setup.sh as the source for expected keys."
                ],
                table=(["Group", "Representative variables", "Source basis"], CONFIG_GROUP_ROWS),
            ),
            Section(
                "Step-By-Step Deployment",
                paragraphs=[
                    "The guided setup path should be used for normal evaluation deployment. Manual commands are appropriate when operators have already prepared .env and want explicit control over build/start timing."
                ],
                table=(
                    ["Step", "Command / action", "Expected result"],
                    [
                        ["1. Review source", "Confirm current commit and inspect README.md, docker-compose.yml, .env.example.", "Operator understands delivery scope and exposed services."],
                        ["2. Prepare .env", "./setup.sh or copy .env.example to .env and replace placeholders.", "Private runtime configuration exists locally and is not committed."],
                        ["3. Validate Compose", "docker compose config --quiet", "Compose interpolation and syntax pass."],
                        ["4. Build images", "docker compose build", "Local service images are built."],
                        ["5. Start stack", "docker compose up -d or setup.sh default start path.", "Containers start on the configured network and volumes."],
                        ["6. Verify", "docker compose ps, health endpoints, verify.sh when services are running.", "Operator confirms live readiness before demo/handover."],
                    ],
                ),
            ),
            Section(
                "Startup, Shutdown, And Clean Reinstall Procedures",
                paragraphs=[
                    "Supported startup paths are ./setup.sh or manual Docker Compose commands after .env is prepared. Supported stop/start procedures use docker compose down/up or docker compose restart without deleting volumes. Destructive volume deletion is not part of routine operations.",
                ],
                table=(
                    ["Procedure", "Command or source", "Notes"],
                    [
                        ["Guided setup", "./setup.sh [--server-ip HOST] [--no-build] [--no-start]", "Creates/updates .env, validates compose, builds and starts by default."],
                        ["Manual validation", "docker compose config", "Base validation succeeded for the current delivery state."],
                        ["Manual start", "docker compose build; docker compose up -d", "Use after .env has no placeholders."],
                        ["Restart service", "docker compose restart analytics", "Use service-specific logs to confirm recovery."],
                        ["Stop without deleting data", "docker compose down", "Keeps persistent volumes."],
                        ["Clean reinstall", "Only when data removal is intended; do not use down -v casually", "Back up data first."],
                    ],
                ),
            ),
            Section(
                "Verification Scripts And Health Checks",
                paragraphs=[
                    "The repository provides verification scripts. These scripts are source material for intended operational checks, but their success depends on a running stack and reachable services. In this documentation run, compose validation was executed; live health checks were not implied."
                ],
                table=(
                    ["Check", "Purpose", "Source/status"],
                    [
                        ["docker compose config --quiet", "Validate Compose syntax/resolution", "Ran successfully for base HumanEnerDIA stack."],
                        ["OVOS-EnMS repository docker-compose.yml", "Validate OVOS assistant Compose configuration separately", "Ran successfully during documentation review."],
                        ["verify.sh", "Checks Compose config and live Nginx, analytics, and optional OVOS endpoints when a stack is running", "Script exists in the production repository; live checks require a running target deployment."],
                        ["Chatbot/Rasa live checks", "Use Compose healthchecks, service logs, and chatbot/Rasa endpoints when services are running", "No production-tracked standalone chatbot verification script is cited."],
                        ["Service healthchecks", "Container-level checks for all production Compose services", "Configured in docker-compose.yml."],
                    ],
                ),
            ),
            Section(
                "Healthcheck Details",
                paragraphs=[
                    "Compose healthchecks are present for the production service inventory. A healthy container does not prove business data quality, but it is the first operational signal for deployment readiness."
                ],
                table=(["Service", "Healthcheck focus", "Interpretation"], HEALTHCHECK_ROWS),
            ),
            Section(
                "Backup And Recovery",
                paragraphs=[
                    "The production tree provides persistent volumes and tracked dashboard/flow configuration, but it does not provide a complete universal backup/restore automation. Operators should implement tested backup procedures before production data is at risk."
                ],
                table=(["Area", "Recommended handling", "Caution"], BACKUP_RECOVERY_ROWS),
            ),
            Section(
                "Upgrade And Redeployment",
                paragraphs=[
                    "Redeployment should be treated as a controlled change. Take backups, rebuild images, restart services, and run smoke checks. If schema/data changes are introduced in future versions, rollback must include database/volume strategy, not only Git checkout."
                ],
                table=(["Phase", "Action"], UPGRADE_REDEPLOY_ROWS),
            ),
            Section(
                "Troubleshooting Commands",
                paragraphs=[
                    "These commands are safe static/live inspection commands when run by an operator with Docker access. They should not delete volumes or modify runtime state."
                ],
                table=(["Purpose", "Command", "Use"], TROUBLESHOOTING_COMMAND_ROWS),
            ),
            Section(
                "Production Hardening Checklist",
                paragraphs=[
                    "The checklist below is intentionally phrased as operator action. The repository provides hooks and configuration, but production hardening is not complete until these actions are performed in the target environment."
                ],
                table=(["Area", "Required action"], HARDENING_ROWS),
            ),
            Section(
                "Troubleshooting Scenarios",
                paragraphs=[
                    "The repository provides deployment-oriented defaults, placeholders, health checks, and hardening notes. Production hardening is completed through operator-managed DNS, TLS, firewall restrictions, credential rotation, backups, and monitoring policy.",
                ],
                table=(
                    ["Symptom", "Likely area", "First checks"],
                    [
                        ["Portal does not load", "Nginx or portal static files", "curl /health; docker compose logs nginx"],
                        ["Analytics API returns 500", "Analytics, PostgreSQL, Redis", "logs for analytics/postgres/redis; /api/v1/health"],
                        ["No new telemetry", "Simulator, MQTT, Node-RED", "logs for simulator/mqtt/nodered; Node-RED flow status"],
                        ["Grafana unavailable", "Grafana or database", "Grafana health endpoint; credentials and volume status"],
                        ["Auth errors", "auth-service, database, SMTP", "/api/auth/health; auth-service logs"],
                        ["OVOS voice path unavailable", "OVOS bridge/messagebus or analytics proxy", "OVOS /health; analytics /api/v1/ovos/voice/health"],
                    ],
                ),
            ),
            common_limitations_section(),
            section_evidence([
                ["Base compose", EVIDENCE["compose"]],
                ["Setup helper", EVIDENCE["setup"]],
                ["Verifier", EVIDENCE["verifier"]],
                ["OVOS Docker", EVIDENCE["ovos_config"]],
            ]),
        ],
    )


def build_final_system(diagrams: dict) -> DocSpec:
    return DocSpec(
        filename="Final System Documentation.docx",
        title="Final System Documentation",
        purpose="Provide a stakeholder-ready end-to-end overview, installation/operation guide, workflows, and final delivery notes.",
        audience="Managers, reviewers, operators, users, integrators, and external WASABI partners.",
        evidence_note="This document summarizes the evidence-backed content of the technical documents and points readers to implemented sources.",
        sections=[
            Section(
                "Project Overview",
                figure=("system-context.png", "Figure 1. Relationship between HumanEnerDIA, OVOS-EnMS, data services, dashboards, and users."),
                paragraphs=[
                    "HumanEnerDIA is an open-source industrial energy management system developed for the WASABI project delivery context. It monitors and analyzes simulated or ingested factory energy data, provides dashboards and APIs, supports ISO 50001-oriented concepts, and integrates with an OVOS assistant layer for natural-language operational queries.",
                    "This final delivery package presents the system as a Docker Compose deployable HumanEnerDIA stack with a companion OVOS-EnMS assistant repository/runtime. It includes implemented backend services, dashboards, reports, simulator, ingestion flow, authentication, text chatbot, and OVOS voice/natural-language paths, with limitations kept visible for review.",
                ],
            ),
            Section(
                "Handover Summary",
                paragraphs=[
                    "The six-document package gives managers, reviewers, operators, and external partners a clear view of what has been delivered, how the parts fit together, how to deploy and verify the stack, and which responsibilities apply to runtime validation and production hardening."
                ],
                table=(["Document", "Primary use"], [
                    ["System Architecture Report", "Architecture, service responsibilities, boundaries, data/message flow, security/network posture, and operational risks."],
                    ["Software Design Documentation", "Module design, API/database/service/auth design, configuration, logging, error handling, and design limitations."],
                    ["Skill Documentation", "OVOS-EnMS query lifecycle, intents, parser/validator/client/formatter behavior, configuration, mappings, and failure behavior."],
                    ["Energy Management System Reports and KPI Reports", "Energy data model, KPI catalog, formulas, dashboards, report workflow, data-quality assumptions, and audit cautions."],
                    ["Docker Deployment Report", "Compose topology, setup, environment groups, health checks, backup/recovery, redeployment, troubleshooting, and hardening."],
                    ["Final System Documentation", "End-to-end reader guide, workflows, operations, acceptance/demo checklists, and final delivery notes."],
                ]),
            ),
            Section(
                "Delivery Artifact List",
                paragraphs=[
                    "The formal deliverables are the six DOCX files. Source Markdown, the evidence map, diagrams, and generator script are included for maintainability and reproducibility."
                ],
                table=(["Artifact", "Purpose"], FINAL_ARTIFACT_ROWS),
            ),
            Section(
                "Reader Guide",
                bullets=[
                    "Managers should start with Final System Documentation, then System Architecture Report, then the limitations sections.",
                    "Technical reviewers should read Software Design Documentation, Energy Management System Reports and KPI Reports, and Docker Deployment Report.",
                    "Assistant/voice reviewers should read Skill Documentation and the OVOS integration sections in the architecture report.",
                    "Operators should use Docker Deployment Report together with setup.sh, verify.sh, docker-compose.yml, and .env.example.",
                    "External partners should treat evidence references as the traceability map for major technical claims.",
                ],
            ),
            Section(
                "Quick-Start Summary",
                paragraphs=[
                    "This quick-start summarizes the handover path. It does not replace the deployment report or the live verification scripts."
                ],
                table=(
                    ["Step", "Action", "Outcome"],
                    [
                        ["1", "Review .env.example and run ./setup.sh with the target server host/IP if needed.", ".env is created locally and first-run secrets are generated."],
                        ["2", "Run or confirm docker compose config --quiet.", "Compose syntax and interpolation are valid."],
                        ["3", "Start the stack through setup.sh or docker compose up -d.", "HumanEnerDIA services start on the configured network and ports."],
                        ["4", "Run docker compose ps and health checks.", "Container status and basic service health are visible."],
                        ["5", "Run verify.sh when services are running.", "Gateway, analytics, and optional OVOS live checks are performed."],
                        ["6", "Open portal, Grafana, analytics UI/API docs, and assistant endpoints as needed.", "Stakeholder/demo access points are ready for review."],
                    ],
                ),
            ),
            Section(
                "Installation And Access",
                paragraphs=[
                    "For a local or evaluation deployment, use the guided setup script from the production repository checkout or approved delivery bundle. For remote browser access, pass a server host or IP so generated URLs match the expected access path. Generated credentials are stored in .env and must be kept private.",
                ],
                table=(["Access point", "Default URL", "Notes"], ACCESS_ROWS),
            ),
            Section(
                "Main Workflows",
                table=(
                    ["Workflow", "Implementation path", "Result"],
                    [
                        ["Start system", "./setup.sh or docker compose build/up", "Services start with generated or configured environment values."],
                        ["Generate telemetry", "simulator -> MQTT -> Node-RED -> PostgreSQL", "Energy, production, environmental, status, and selected multi-energy data are stored."],
                        ["View dashboards", "Grafana through Nginx or direct port", "Configured SOTA dashboard JSON views are available."],
                        ["Use analytics APIs", "FastAPI analytics service under /api/v1", "KPIs, baselines, forecasts, anomalies, reports, and related data are exposed."],
                        ["Authenticate/admin", "auth-service through portal and /api/auth, /api/admin", "Registration/login/admin/session flows use JWT and database tables."],
                        ["Ask text help questions", "chatbot Express proxy -> Rasa -> custom action QA retrieval", "Knowledge/help answers from qa_data.json categories."],
                        ["Ask operational assistant questions", "OVOS REST bridge -> messagebus -> EnMS skill -> analytics API", "Voice/text operational responses with structured data."],
                        ["Generate reports", "analytics report endpoints", "Legacy monthly EnPI PDF and V2 PDF paths are available."],
                    ],
                ),
            ),
            Section(
                "Operator Guide",
                bullets=[
                    "Use docker compose ps and service health endpoints for daily status checks.",
                    "Check Nginx first for browser routing issues, then the owning upstream service.",
                    "Inspect simulator, MQTT, Node-RED, and PostgreSQL together for data-ingestion issues.",
                    "Use the tracked JSON files under grafana/dashboards and Grafana provisioning files for dashboard review and controlled updates.",
                    "Use pg_dump or platform backup tooling for PostgreSQL; no generic tracked database backup script exists.",
                    "Avoid docker compose down -v unless the purpose is deliberate persistent data deletion.",
                ],
            ),
            Section(
                "Analytics, Dashboards, Reports, And Assistants",
                paragraphs=[
                    "Analytics capabilities include baseline training/prediction/deviation, KPI functions, forecasting, anomaly detection/search, machine status and time series, comparison/visualization data, model performance, production metrics, performance analysis, ISO 50001/SEU endpoints, and report generation.",
                    "Dashboards are configured in Grafana JSON and provisioned through the Grafana provisioning directory. The Rasa chatbot is a text help/knowledge assistant, while OVOS-EnMS is the operational natural-language assistant integrated with live backend APIs.",
                ],
                table=(["Dashboard or assistant", "Purpose"], [
                    ["Grafana dashboards", "Operational, executive, cost, carbon, ISO 50001, anomaly, model, production, and predictive views."],
                    ["Analytics UI", "FastAPI-rendered pages for dashboards, baselines, anomalies, KPIs, forecasts, Sankey, heatmap, comparison, and model performance."],
                    ["Rasa chatbot", "Text knowledge/help assistant using QA categories and custom retrieval action."],
                    ["OVOS-EnMS", "Operational assistant for energy, power, machine status, rankings, anomalies, forecasts, baselines, KPIs, reports, help, and health checks."],
                ]),
            ),
            Section(
                "Maintenance And Troubleshooting",
                paragraphs=[
                    "Maintenance should focus on credential rotation, backup verification, dashboard export/commit policy, disk usage, restart counts, recent errors, image/base dependency review, and firewall/public route review. Production hardening requires operator policy beyond the repository defaults.",
                ],
                table=(
                    ["Maintenance area", "Recommended review"],
                    [
                        ["Credentials", "Rotate generated first-run values and any exposed credentials."],
                        ["Backups", "Test PostgreSQL restore; back up Grafana dashboards and Docker volumes."],
                        ["Dashboards", "Commit intended Grafana JSON changes after backup/export."],
                        ["Telemetry", "Confirm simulator/MQTT/Node-RED/PostgreSQL are all healthy when data appears stale."],
                        ["OVOS", "Confirm /health messagebus_connected and smoke query when assistant is required."],
                        ["Docs", "Update final documents when route, schema, compose, or packaging behavior changes."],
                    ],
                ),
            ),
            Section(
                "Final Delivery Notes",
                bullets=[
                    "The documentation package relies on the GitHub production source tree and the separate OVOS-EnMS repository, not on untracked local release artifacts.",
                    ".env is not shipped and must not be disclosed.",
                    "The optional Qwen GGUF model path and LLM fallback controls are documented from the OVOS-EnMS code/config; model availability must be verified in the deployed OVOS-EnMS runtime.",
                    "Runtime health verification belongs to the actual target deployment used for demonstration or handover.",
                ],
            ),
            Section(
                "Acceptance Checklist",
                paragraphs=[
                    "The following checks summarize the deliverable acceptance state for the documentation set."
                ],
                table=(["Check", "Expected result"], ACCEPTANCE_CHECK_ROWS),
            ),
            Section(
                "Demo-Readiness Checklist",
                paragraphs=[
                    "For a live demonstration, use this checklist alongside the deployment report and verify.sh. These checks are operational and must be performed on the actual target deployment."
                ],
                table=(["Area", "Readiness action"], DEMO_READINESS_ROWS),
            ),
            common_limitations_section(),
            section_evidence([
                ["Overview/docs", f"{EVIDENCE['readme']}; docs/final-delivery/"],
                ["Deployment", f"{EVIDENCE['compose']}; {EVIDENCE['setup']}; {EVIDENCE['verifier']}"],
                ["Analytics/database", f"{EVIDENCE['analytics_main']}; {EVIDENCE['database_schema']}"],
                ["Dashboards and reports", f"{EVIDENCE['grafana']}; {EVIDENCE['reports']}"],
                ["OVOS-EnMS", "OVOS-EnMS repository: README.md; enms-ovos-skill/README.md; enms-ovos-skill/bridge/ovos_rest_bridge.py; enms-ovos-skill/enms_ovos_skill/__init__.py"],
                ["Delivery package sources", "docs/final-delivery/source/; docs/final-delivery/assets/"],
            ]),
        ],
    )


def build_markdown(spec: DocSpec) -> str:
    lines = [
        f"# {spec.title}",
        "",
        f"Project: {PROJECT_NAME}",
        f"Version: {DOC_VERSION}",
        f"Date: {DOC_DATE}",
        f"Status: {DOC_STATUS}",
        "",
        f"Purpose: {spec.purpose}",
        f"Audience: {spec.audience}",
        "",
        f"Source basis: {spec.evidence_note}",
        "",
    ]
    for section in spec.sections:
        render_markdown_section(lines, section, 2)
    return "\n".join(lines).rstrip() + "\n"


def render_markdown_section(lines: List[str], section: Section, level: int) -> None:
    lines.append(f"{'#' * level} {section.title}")
    lines.append("")
    if section.figure:
        image_name, caption = section.figure
        lines.append(f"![{caption}](../assets/{image_name})")
        lines.append("")
    for para in section.paragraphs:
        if isinstance(para, tuple):
            lines.append(f"**{para[0].rstrip()}** {para[1].lstrip()}")
        else:
            lines.append(para)
        lines.append("")
    for bullet in section.bullets:
        lines.append(f"- {bullet}")
    if section.bullets:
        lines.append("")
    if section.table:
        intro = table_intro_for(section.title)
        if intro:
            if isinstance(intro, tuple):
                lines.append(f"**{intro[0].rstrip()}** {intro[1].lstrip()}")
            else:
                lines.append(intro)
            lines.append("")
        headers, rows = section.table
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            lines.append("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |")
        lines.append("")
    followup_blocks = section.table_followup or ([table_followup_for(section.title)] if section.table else [])
    for para in followup_blocks:
        if isinstance(para, tuple):
            lines.append(f"**{para[0].rstrip()}** {para[1].lstrip()}")
        else:
            lines.append(para)
        lines.append("")
    for sub in section.subsections:
        render_markdown_section(lines, sub, level + 1)


def write_sources(specs: Sequence[DocSpec]) -> None:
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    for spec in specs:
        source_name = spec.filename.replace(".docx", ".md")
        (SOURCE_DIR / source_name).write_text(build_markdown(spec), encoding="utf-8")

    evidence_lines = [
        "# Evidence Map",
        "",
        f"Project: {PROJECT_NAME}",
        f"Generated: {DOC_DATE}",
        "",
        "This file maps recurring documentation claims to local evidence. It intentionally avoids .env values and credential-bearing runtime files.",
        "",
        "| Key | Evidence |",
        "| --- | --- |",
    ]
    for key in sorted(EVIDENCE):
        evidence_lines.append(f"| {key} | {EVIDENCE[key]} |")
    evidence_lines.extend([
        "",
        "## Validation Performed",
        "",
        "- `docker compose config --quiet` in the HumanEnerDIA production tree: passed.",
        "- `docker compose -f <OVOS-EnMS repository>/docker-compose.yml config --quiet`: passed.",
        "- Runtime health checks were not run by this generator; they require a running deployment.",
    ])
    (SOURCE_DIR / "evidence-map.md").write_text("\n".join(evidence_lines) + "\n", encoding="utf-8")


def build_specs(diagrams: dict) -> List[DocSpec]:
    return [
        build_system_architecture(diagrams),
        build_software_design(),
        build_skill_doc(diagrams),
        build_kpi_doc(diagrams),
        build_docker_doc(diagrams),
        build_final_system(diagrams),
    ]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = generate_diagrams()
    specs = build_specs(diagrams)
    write_sources(specs)
    for spec in specs:
        build_document(spec)
    print("Generated DOCX files:")
    for spec in specs:
        print(OUT_DIR / spec.filename)


if __name__ == "__main__":
    main()
